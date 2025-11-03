import sys
import os
import tempfile
import time
import sqlite3
import json
import shutil
from datetime import datetime
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                              QHBoxLayout, QPushButton, QLabel, QTextEdit, 
                              QFileDialog, QProgressBar, QGroupBox, QSpinBox,
                              QDoubleSpinBox, QCheckBox, QComboBox, QMessageBox,
                              QProgressDialog, QTabWidget, QListWidget, QListWidgetItem,
                              QSplitter, QScrollArea, QSizePolicy)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QSize, QRect, QPoint
from PyQt6.QtGui import QPixmap, QImage, QPainter, QPen, QColor
# Lazy import torch and transformers để tăng tốc khởi động
# Chỉ import khi cần load model
# import torch
# from transformers import Qwen3VLForConditionalGeneration, AutoProcessor
from PIL import Image

# Lazy import functions để tăng tốc khởi động
_torch = None
_transformers = None

def get_torch():
    """Lazy import torch - chỉ import khi cần"""
    global _torch
    if _torch is None:
        import torch
        _torch = torch
    return _torch

def get_transformers():
    """Lazy import transformers - chỉ import khi cần"""
    global _transformers
    if _transformers is None:
        from transformers import Qwen3VLForConditionalGeneration, AutoProcessor
        _transformers = {
            'Qwen3VLForConditionalGeneration': Qwen3VLForConditionalGeneration,
            'AutoProcessor': AutoProcessor
        }
    return _transformers

def is_cuda_available():
    """Check CUDA availability - lazy check"""
    try:
        torch = get_torch()
        return torch.cuda.is_available()
    except:
        return False

def get_cuda_info():
    """Get CUDA info - lazy check"""
    try:
        torch = get_torch()
        if torch.cuda.is_available():
            return {
                'available': True,
                'device_name': torch.cuda.get_device_name(0),
                'cuda_version': torch.version.cuda,
                'gpu_memory': torch.cuda.get_device_properties(0).total_memory / (1024**3)
            }
    except:
        pass
    return {'available': False}

# System check
try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False
    print("Cảnh báo: psutil chưa được cài đặt. Kiểm tra RAM bị tắt. Cài đặt: pip install psutil")

# PDF handling
try:
    from pdf2image import convert_from_path
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("Cảnh báo: pdf2image chưa được cài đặt. Hỗ trợ PDF bị tắt. Cài đặt: pip install pdf2image")

# DOCX handling
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("Cảnh báo: python-docx chưa được cài đặt. Hỗ trợ DOCX bị tắt. Cài đặt: pip install python-docx")


class ImageROILabel(QLabel):
    """Custom QLabel for image display with ROI selection"""
    roi_changed = pyqtSignal(QRect)  # Emit when ROI changes
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.original_pixmap = None
        self.current_pixmap = None
        self.roi_mode = False
        self.roi_rect = QRect()
        self.start_point = QPoint()
        self.end_point = QPoint()
        self.is_drawing = False
        self.setMouseTracking(True)
    
    def set_pixmap(self, pixmap):
        """Set image pixmap"""
        self.original_pixmap = pixmap
        self.current_pixmap = pixmap
        self.update_display()
    
    def set_roi_mode(self, enabled):
        """Enable/disable ROI mode"""
        self.roi_mode = enabled
        if not enabled:
            self.roi_rect = QRect()
            self.start_point = QPoint()
            self.end_point = QPoint()
            self.is_drawing = False
        self.update_display()
    
    def get_roi_rect(self):
        """Get ROI rectangle in image coordinates"""
        if self.roi_rect.isEmpty():
            return None
        return self.roi_rect
    
    def clear_roi(self):
        """Clear ROI selection"""
        self.roi_rect = QRect()
        self.update_display()
        self.roi_changed.emit(QRect())
    
    def screen_to_image_coords(self, screen_point):
        """Convert screen coordinates to image coordinates"""
        if not self.original_pixmap or self.original_pixmap.isNull():
            return QPoint()
        
        # Get label size
        label_size = self.size()
        if label_size.width() <= 0 or label_size.height() <= 0:
            return QPoint()
        
        pixmap_size = self.original_pixmap.size()
        
        # Calculate scaled pixmap size maintaining aspect ratio
        scaled_size = pixmap_size.scaled(label_size, Qt.AspectRatioMode.KeepAspectRatio)
        
        # Calculate offset (centered)
        offset_x = (label_size.width() - scaled_size.width()) // 2
        offset_y = (label_size.height() - scaled_size.height()) // 2
        
        # Convert to image coordinates
        image_x = int((screen_point.x() - offset_x) * pixmap_size.width() / scaled_size.width())
        image_y = int((screen_point.y() - offset_y) * pixmap_size.height() / scaled_size.height())
        
        # Clamp to image bounds
        image_x = max(0, min(image_x, pixmap_size.width() - 1))
        image_y = max(0, min(image_y, pixmap_size.height() - 1))
        
        return QPoint(image_x, image_y)
    
    def mousePressEvent(self, event):
        """Handle mouse press"""
        if self.roi_mode and event.button() == Qt.MouseButton.LeftButton:
            self.is_drawing = True
            self.start_point = self.screen_to_image_coords(event.position().toPoint())
            self.end_point = self.start_point
            self.update_display()
        super().mousePressEvent(event)
    
    def mouseMoveEvent(self, event):
        """Handle mouse move"""
        if self.roi_mode and self.is_drawing:
            self.end_point = self.screen_to_image_coords(event.position().toPoint())
            self.update_display()
        super().mouseMoveEvent(event)
    
    def mouseReleaseEvent(self, event):
        """Handle mouse release"""
        if self.roi_mode and event.button() == Qt.MouseButton.LeftButton and self.is_drawing:
            self.is_drawing = False
            self.end_point = self.screen_to_image_coords(event.position().toPoint())
            self.update_roi_rect()
            self.update_display()
            self.roi_changed.emit(self.roi_rect)
        super().mouseReleaseEvent(event)
    
    def update_roi_rect(self):
        """Update ROI rectangle from start and end points"""
        if not self.original_pixmap or self.original_pixmap.isNull():
            return
        
        pixmap_size = self.original_pixmap.size()
        
        # Normalize rectangle
        x1 = min(self.start_point.x(), self.end_point.x())
        y1 = min(self.start_point.y(), self.end_point.y())
        x2 = max(self.start_point.x(), self.end_point.x())
        y2 = max(self.start_point.y(), self.end_point.y())
        
        # Clamp to image bounds
        x1 = max(0, min(x1, pixmap_size.width() - 1))
        y1 = max(0, min(y1, pixmap_size.height() - 1))
        x2 = max(0, min(x2, pixmap_size.width() - 1))
        y2 = max(0, min(y2, pixmap_size.height() - 1))
        
        self.roi_rect = QRect(QPoint(x1, y1), QPoint(x2, y2))
    
    def update_display(self):
        """Update displayed image with ROI overlay"""
        if not self.original_pixmap or self.original_pixmap.isNull():
            return
        
        # Scale pixmap to fit label
        label_size = self.size()
        if label_size.width() <= 0 or label_size.height() <= 0:
            return
        
        scaled_pixmap = self.original_pixmap.scaled(
            label_size,
            Qt.AspectRatioMode.KeepAspectRatio,
            Qt.TransformationMode.SmoothTransformation
        )
        
        # Draw ROI rectangle if in ROI mode
        if self.roi_mode and (not self.roi_rect.isEmpty() or self.is_drawing):
            # Create painter on pixmap
            painter = QPainter(scaled_pixmap)
            painter.setRenderHint(QPainter.RenderHint.Antialiasing)
            
            # Calculate ROI in scaled coordinates
            pixmap_size = self.original_pixmap.size()
            scaled_size = scaled_pixmap.size()
            scale_x = scaled_size.width() / pixmap_size.width()
            scale_y = scaled_size.height() / pixmap_size.height()
            
            if self.is_drawing:
                # Draw temporary rectangle while drawing
                x1 = min(self.start_point.x(), self.end_point.x()) * scale_x
                y1 = min(self.start_point.y(), self.end_point.y()) * scale_y
                x2 = max(self.start_point.x(), self.end_point.x()) * scale_x
                y2 = max(self.start_point.y(), self.end_point.y()) * scale_y
                roi_scaled = QRect(int(x1), int(y1), int(x2 - x1), int(y2 - y1))
            else:
                # Draw final ROI rectangle
                roi_scaled = QRect(
                    int(self.roi_rect.x() * scale_x),
                    int(self.roi_rect.y() * scale_y),
                    int(self.roi_rect.width() * scale_x),
                    int(self.roi_rect.height() * scale_y)
                )
            
            # Draw semi-transparent overlay
            overlay_color = QColor(0, 123, 255, 100)  # Blue with transparency
            painter.fillRect(scaled_pixmap.rect(), overlay_color)
            
            # Draw ROI rectangle with border
            pen = QPen(QColor(0, 123, 255), 3)
            painter.setPen(pen)
            painter.drawRect(roi_scaled)
            
            # Clear inside ROI
            painter.setCompositionMode(QPainter.CompositionMode.CompositionMode_Clear)
            painter.fillRect(roi_scaled, QColor(0, 0, 0, 0))
            
            painter.end()
        
        # Set scaled pixmap
        self.current_pixmap = scaled_pixmap
        self.setPixmap(scaled_pixmap)
    
    def resizeEvent(self, event):
        """Handle resize event"""
        super().resizeEvent(event)
        if self.original_pixmap:
            self.update_display()


class ModelLoaderWorker(QThread):
    """Worker thread for loading model"""
    finished = pyqtSignal(object, object)  # model, processor
    error = pyqtSignal(str)
    progress = pyqtSignal(str)
    
    def __init__(self, model_path, device_map, dtype, low_cpu_mem_usage, max_memory):
        super().__init__()
        self.model_path = model_path
        self.device_map = device_map
        self.dtype = dtype
        self.low_cpu_mem_usage = low_cpu_mem_usage
        self.max_memory = max_memory
        self.model = None
        self.processor = None
    
    def run(self):
        """Load model in worker thread với comprehensive error handling"""
        import traceback
        
        try:
            print("[Worker Thread] Bắt đầu load model...")
            self.progress.emit("Đang khởi tạo model...")
            
            # Load model with comprehensive error handling
            try:
                print(f"[Worker Thread] Đang load model từ: {self.model_path}")
                print(f"[Worker Thread] Device: {self.device_map}, dtype: {self.dtype}")
                
                # Check available memory before loading
                try:
                    import psutil
                    ram = psutil.virtual_memory()
                    ram_available_gb = ram.available / (1024**3)
                    print(f"[Worker Thread] RAM khả dụng: {ram_available_gb:.1f}GB")
                    
                    if ram_available_gb < 4:
                        raise Exception(
                            f"RAM không đủ! Chỉ có {ram_available_gb:.1f}GB khả dụng.\n"
                            f"Cần ít nhất 4GB RAM trống để load model.\n\n"
                            f"Giải pháp:\n"
                            f"1. Đóng các ứng dụng khác\n"
                            f"2. Khởi động lại máy\n"
                            f"3. Tăng paging file lên ít nhất 8GB"
                        )
                except ImportError:
                    print("[Worker Thread] Không thể kiểm tra RAM (psutil không có)")
                except Exception as e:
                    error_msg = str(e)
                    if "RAM không đủ" in error_msg:
                        raise Exception(error_msg)
                
                # Load model với safer options
                print("[Worker Thread] Bắt đầu load checkpoint shards...")
                print("[Worker Thread] LƯU Ý: Quá trình này có thể mất 2-5 phút và có thể crash nếu memory không đủ!")
                print("[Worker Thread] Nếu crash, vui lòng tăng paging file lên ít nhất 8GB và restart máy.")
                
                # Flush output để đảm bảo log được ghi trước khi crash
                import sys
                sys.stdout.flush()
                sys.stderr.flush()
                
                # Emit progress để UI biết đang ở đâu
                self.progress.emit("Đang load checkpoint shards... (Có thể mất 2-5 phút)")
                
                # Wrap trong try-except để catch mọi exception
                try:
                    # Import ngay trong try để catch import errors
                    from transformers import Qwen3VLForConditionalGeneration
                    
                    # Load model - có thể mất 2-5 phút
                    # QUAN TRỌNG: Nếu crash xảy ra ở đây (C++ level), Python không thể catch!
                    # Crash có thể do: out of memory kill by Windows, segmentation fault, etc.
                    print("[Worker Thread] Nếu crash ở đây, có nghĩa là memory không đủ hoặc paging file quá nhỏ!")
                    sys.stdout.flush()
                    
                    # Thử load với options đơn giản như code mẫu Hugging Face
                    # Code mẫu dùng dtype="auto" và device_map="auto"
                    print(f"[Worker Thread] Loading với device_map='{self.device_map}', dtype={self.dtype}")
                    
                    # Kiểm tra paging file size - ĐÂY LÀ ĐIỂM QUAN TRỌNG NHẤT!
                    try:
                        import psutil
                        swap = psutil.swap_memory()
                        swap_total_gb = swap.total / (1024**3)
                        swap_used_gb = swap.used / (1024**3)
                        
                        # Tính toán paging file cần thiết dựa trên thiết bị đã chọn
                        ram = psutil.virtual_memory()
                        ram_total_gb = ram.total / (1024**3)
                        
                        # Tính model size dựa trên model được chọn (2B hoặc 4B)
                        # Từ model_path có thể suy ra model size
                        if "2B" in self.model_path:
                            model_size_gb = 4.0  # Model 2B float16 ≈ 4GB
                            model_name = "2B"
                        else:
                            model_size_gb = 8.0  # Model 4B float16 ≈ 8GB
                            model_name = "4B"
                        
                        torch = get_torch()
                        if self.device_map == "cuda:0" and is_cuda_available():
                            # GPU mode: Model sẽ load vào VRAM, paging file chỉ cần model size cho loading process
                            gpu_memory_gb = torch.cuda.get_device_properties(0).total_memory / (1024**3)
                            # Công thức: Model size (VRAM đủ cho model, paging chỉ cần cho loading)
                            paging_needed_gb = model_size_gb
                            print(f"[Worker Thread] GPU mode - VRAM: {gpu_memory_gb:.1f}GB | Model {model_name}: {model_size_gb}GB | Paging cần: {paging_needed_gb}GB")
                        else:
                            # CPU mode: Model load vào RAM, paging file cần model size
                            # Công thức: Model size (CPU mode cần paging = model size)
                            paging_needed_gb = model_size_gb
                            print(f"[Worker Thread] CPU mode - RAM: {ram_total_gb:.1f}GB | Model {model_name}: {model_size_gb}GB | Paging cần: {paging_needed_gb}GB")
                        
                        # Làm tròn lên đến GB gần nhất (tối thiểu = model size)
                        paging_needed_gb = max(int(model_size_gb), int(paging_needed_gb))
                        print(f"[Worker Thread] Paging file: {swap_used_gb:.1f}GB / {swap_total_gb:.1f}GB")
                        
                        torch = get_torch()
                        if self.device_map == "cuda:0" and is_cuda_available():
                            gpu_memory_gb = torch.cuda.get_device_properties(0).total_memory / (1024**3)
                            print(f"[Worker Thread] Paging file cần thiết: ~{paging_needed_gb}GB (dựa trên GPU VRAM: {gpu_memory_gb:.1f}GB)")
                        else:
                            print(f"[Worker Thread] Paging file cần thiết: ~{paging_needed_gb}GB (dựa trên RAM: {ram_total_gb:.1f}GB)")
                        
                        # Kiểm tra nếu paging file quá nhỏ (ít hơn 80% của cần thiết)
                        if swap_total_gb < paging_needed_gb * 0.8:
                            # PAGING FILE QUÁ NHỎ - DỪNG NGAY VÀ BÁO LỖI!
                            torch = get_torch()
                            if self.device_map == "cuda:0" and is_cuda_available():
                                gpu_memory_gb = torch.cuda.get_device_properties(0).total_memory / (1024**3)
                                device_info = f"GPU (VRAM: {gpu_memory_gb:.1f}GB)"
                                memory_info = f"GPU VRAM: {gpu_memory_gb:.1f}GB"
                                formula_detail = f"Model size ({model_name} float16): {model_size_gb}GB = {paging_needed_gb:.0f}GB"
                            else:
                                device_info = "CPU"
                                memory_info = f"RAM CPU: {ram_total_gb:.1f}GB"
                                formula_detail = f"Model size ({model_name} float16): {model_size_gb}GB = {paging_needed_gb:.0f}GB"
                            
                            error_msg = (
                                f"PAGING FILE QUÁ NHỎ!\n\n"
                                f"Thông tin hệ thống:\n"
                                f"- Thiết bị đã chọn: {device_info}\n"
                                f"- {memory_info}\n"
                                f"- Paging file hiện tại: {swap_total_gb:.1f}GB\n"
                                f"- Paging file cần thiết: ~{paging_needed_gb}GB\n\n"
                                f"Paging file quá nhỏ sẽ khiến ứng dụng crash khi load model!\n\n"
                                f"GIẢI PHÁP (LÀM THEO THỨ TỰ):\n\n"
                                f"1. Mở System Properties:\n"
                                f"   - Nhấn Win+R\n"
                                f"   - Gõ: sysdm.cpl\n"
                                f"   - Nhấn Enter\n\n"
                                f"2. Tăng Virtual Memory:\n"
                                f"   - Tab 'Advanced'\n"
                                f"   - Click 'Settings' trong Performance\n"
                                f"   - Tab 'Advanced'\n"
                                f"   - Click 'Change' trong Virtual memory\n"
                                f"   - Bỏ tick 'Automatically manage paging file size'\n"
                                f"   - Chọn ổ C:\n"
                                f"   - Chọn 'Custom size'\n"
                                f"   - Initial size: {int(paging_needed_gb * 1024)} (MB) ← {paging_needed_gb:.0f}GB\n"
                                f"   - Maximum size: {int(paging_needed_gb * 1024)} (MB) ← {paging_needed_gb:.0f}GB\n"
                                f"   - Click 'Set'\n"
                                f"   - Click 'OK'\n"
                                f"   - Click 'OK' lần nữa\n\n"
                                f"3. RESTART MÁY (BẮT BUỘC!)\n\n"
                                f"4. Sau khi restart, chạy lại ứng dụng\n\n"
                            )
                            torch = get_torch()
                            device_mode = 'GPU mode' if self.device_map == 'cuda:0' and is_cuda_available() else 'CPU mode'
                            error_msg += (
                                f"Công thức tính ({device_mode}):\n"
                                f"   {formula_detail}\n"
                                f"   = ~{paging_needed_gb}GB\n\n"
                                f"KHÔNG THỂ LOAD MODEL nếu paging file < {paging_needed_gb}GB!\n"
                                f"Đây là giới hạn của Windows khi load model lớn."
                            )
                            
                            print(f"\n{'='*80}")
                            print("[Worker Thread] PHÁT HIỆN PAGING FILE QUÁ NHỎ!")
                            print(f"[Worker Thread] Paging file: {swap_total_gb:.1f}GB")
                            print(f"[Worker Thread] Cần: ít nhất {paging_needed_gb}GB")
                            print(f"{'='*80}\n")
                            
                            # Emit error signal ngay
                            self.error.emit(error_msg)
                            return  # Dừng ngay, không load model
                    except Exception as check_error:
                        print(f"[Worker Thread] Không thể kiểm tra paging file: {check_error}")
                        # Tiếp tục load nếu không thể kiểm tra
                    
                    # Lazy import torch and transformers
                    torch = get_torch()
                    transformers_dict = get_transformers()
                    Qwen3VLForConditionalGeneration = transformers_dict['Qwen3VLForConditionalGeneration']
                    
                    # Với GPU, FORCE load vào GPU VRAM - không dùng auto
                    if self.device_map == "cuda:0" and is_cuda_available():
                        print("[Worker Thread] GPU mode - FORCE loading vào GPU VRAM...")
                        print(f"[Worker Thread] Device map: {self.device_map}, dtype: {self.dtype}")
                        try:
                            # FORCE load vào GPU - dùng device_map="cuda:0" thay vì "auto"
                            # "auto" có thể gây confusion hoặc không load vào GPU đúng cách
                            self.model = Qwen3VLForConditionalGeneration.from_pretrained(
                                self.model_path,
                                dtype=self.dtype if self.dtype is not None else torch.float16,  # GPU thường dùng float16
                                device_map="cuda:0",  # FORCE GPU - rõ ràng chỉ định GPU
                                low_cpu_mem_usage=self.low_cpu_mem_usage,
                                max_memory=self.max_memory,
                                offload_folder=None,
                            )
                            # Sau khi load, đảm bảo model ở GPU
                            if hasattr(self.model, 'device'):
                                print(f"[Worker Thread] Model device: {self.model.device}")
                            else:
                                actual_device = next(self.model.parameters()).device
                                print(f"[Worker Thread] Model device: {actual_device}")
                            print("[Worker Thread] Model loaded vào GPU VRAM thành công!")
                        except Exception as gpu_load_error:
                            error_str = str(gpu_load_error)
                            print(f"[Worker Thread] Load vào GPU thất bại: {error_str}")
                            print("[Worker Thread] Có thể GPU VRAM không đủ, thử load với options tối ưu hơn...")
                            try:
                                # Fallback với options tối ưu memory
                                self.model = Qwen3VLForConditionalGeneration.from_pretrained(
                                    self.model_path,
                                    dtype=self.dtype if self.dtype is not None else torch.float16,
                                    device_map="cuda:0",  # Vẫn force GPU
                                    low_cpu_mem_usage=True,
                                    max_memory=self.max_memory if self.max_memory else {0: "10GiB"},
                                    offload_folder=None,
                                )
                                # Đảm bảo model ở GPU
                                actual_device = next(self.model.parameters()).device
                                print(f"[Worker Thread] Model device: {actual_device}")
                                print("[Worker Thread] Model loaded với options tối ưu memory thành công!")
                            except Exception as fallback_error:
                                # Nếu vẫn thất bại, re-raise để báo lỗi
                                raise Exception(f"Không thể load model vào GPU: {fallback_error}")
                    else:
                        # CPU mode - FORCE load vào CPU, không cho phép auto chọn GPU
                        print("[Worker Thread] CPU mode - FORCE loading vào CPU RAM...")
                        print(f"[Worker Thread] Device map: {self.device_map}, dtype: {self.dtype}")
                        
                        # FORCE load vào CPU - không dùng device_map="auto" vì nó sẽ tự chọn GPU
                        # Phải dùng device_map="cpu" để force load vào CPU RAM
                        self.model = Qwen3VLForConditionalGeneration.from_pretrained(
                            self.model_path,
                            dtype=self.dtype if self.dtype is not None else torch.float32,  # CPU thường dùng float32
                            device_map="cpu",  # FORCE CPU - không cho phép auto chọn GPU
                            low_cpu_mem_usage=self.low_cpu_mem_usage,
                            max_memory=None,  # CPU mode không dùng max_memory
                            offload_folder=None,
                        )
                        # Sau khi load, đảm bảo model ở CPU
                        self.model = self.model.to("cpu")
                        print("[Worker Thread] Model loaded vào CPU RAM thành công!")
                        print(f"[Worker Thread] Model device: {next(self.model.parameters()).device}")
                    print("[Worker Thread] Model loaded successfully!")
                    sys.stdout.flush()
                    
                except KeyboardInterrupt:
                    raise  # Re-raise keyboard interrupt
                except SystemExit:
                    raise  # Re-raise system exit
                except BaseException as load_error:
                    # Catch ALL exceptions including fatal ones
                    error_str = str(load_error)
                    error_type = type(load_error).__name__
                    
                    # Flush output ngay lập tức
                    import sys
                    sys.stdout.flush()
                    sys.stderr.flush()
                    
                    print(f"\n{'='*80}")
                    print(f"[Worker Thread] CRITICAL ERROR DURING LOAD: {error_type}")
                    print(f"[Worker Thread] Error: {error_str}")
                    print(f"{'='*80}")
                    traceback.print_exc()
                    sys.stdout.flush()
                    
                    # Emit error signal - ĐÂY LÀ CÁCH DUY NHẤT để main thread biết có lỗi
                    detailed_error = (
                        f"Lỗi nghiêm trọng khi load checkpoint shards!\n\n"
                        f"Loại lỗi: {error_type}\n"
                        f"Thông báo: {error_str}\n\n"
                        f"LƯU Ý: Nếu ứng dụng CRASH mà không hiển thị popup này,\n"
                        f"có nghĩa là crash xảy ra ở C++ level (PyTorch) mà Python không catch được!\n\n"
                        f"Nguyên nhân có thể:\n"
                        f"1. Windows kill process do out of memory (segmentation fault)\n"
                        f"2. Paging file quá nhỏ (< 8GB)\n"
                        f"3. RAM không đủ (< 16GB total, < 8GB available)\n"
                        f"4. Model files bị corrupt\n"
                        f"5. CUDA/GPU driver issues\n\n"
                        f"GIẢI PHÁP (LÀM THEO THỨ TỰ):\n"
                        f"1. Tăng Windows paging file lên ít nhất 8GB:\n"
                        f"   - Win+R → sysdm.cpl → Advanced → Virtual memory → Custom size: 8192MB\n"
                        f"2. RESTART MÁY (quan trọng!)\n"
                        f"3. Đóng TẤT CẢ ứng dụng khác\n"
                        f"4. Thử chế độ CPU thay vì GPU (chọn từ dropdown)\n"
                        f"5. Check model files có đầy đủ không\n\n"
                        f"Full traceback:\n{traceback.format_exc()}"
                    )
                    
                    # Emit error signal NGAY LẬP TỨC
                    try:
                        self.error.emit(detailed_error)
                        print("[Worker Thread] Error signal emitted successfully!")
                        sys.stdout.flush()
                    except Exception as emit_error:
                        print(f"[Worker Thread] Failed to emit error signal: {emit_error}")
                        sys.stdout.flush()
                    
                    return  # Don't re-raise, already emitted error signal
                
            except RuntimeError as e:
                error_str = str(e)
                print(f"[Worker Thread] RuntimeError: {error_str}")
                traceback.print_exc()
                
                # Check for memory-related errors
                memory_keywords = [
                    "out of memory", "paging file", "1455", "memory", 
                    "cuda out of memory", "alloc", "cannot allocate"
                ]
                is_memory_error = any(keyword in error_str.lower() for keyword in memory_keywords)
                
                if is_memory_error:
                    detailed_error = (
                        f"Lỗi Memory khi load model!\n\n"
                        f"Lỗi gốc: {error_str}\n\n"
                        f"Giải pháp:\n"
                        f"1. Tăng Windows paging file lên ít nhất 8GB (chạy: increase_paging_file.bat)\n"
                        f"2. Đóng tất cả ứng dụng khác để giải phóng RAM\n"
                        f"3. Khởi động lại máy\n"
                        f"4. Thử chế độ CPU thay vì GPU (chọn CPU từ dropdown)\n"
                        f"5. Hệ thống cần ít nhất 16GB RAM (khuyến nghị 32GB)\n\n"
                        f"Traceback:\n{traceback.format_exc()}"
                    )
                    self.error.emit(detailed_error)
                    return
                else:
                    # Other RuntimeError - re-raise để được catch bên ngoài
                    raise
            except Exception as e:
                error_str = str(e)
                error_type = type(e).__name__
                print(f"[Worker Thread] Exception ({error_type}) khi load model: {error_str}")
                traceback.print_exc()
                
                detailed_error = (
                    f"Lỗi khi load model!\n\n"
                    f"Loại lỗi: {error_type}\n"
                    f"Thông báo: {error_str}\n\n"
                    f"Full traceback:\n{traceback.format_exc()}"
                )
                self.error.emit(detailed_error)
                return
            
            self.progress.emit("Đang tải processor...")
            print("[Worker Thread] Đang load processor...")
            
            # Load processor
            try:
                transformers_dict = get_transformers()
                AutoProcessor = transformers_dict['AutoProcessor']
                self.processor = AutoProcessor.from_pretrained(self.model_path)
                print("[Worker Thread] Processor loaded successfully!")
            except Exception as e:
                error_str = str(e)
                print(f"[Worker Thread] Exception khi load processor: {error_str}")
                traceback.print_exc()
                self.error.emit(f"Lỗi khi load processor: {error_str}\n\nTraceback:\n{traceback.format_exc()}")
                return
            
            self.progress.emit("Đang hoàn tất...")
            print("[Worker Thread] Đang clear cache...")
            
            # Clear cache
            torch = get_torch()
            if is_cuda_available():
                torch.cuda.empty_cache()
            
            print("[Worker Thread] Emitting finished signal...")
            self.finished.emit(self.model, self.processor)
            print("[Worker Thread] Finished signal emitted!")
            
        except KeyboardInterrupt:
            print("[Worker Thread] Interrupted by user")
            self.error.emit("Tải model bị hủy bởi người dùng")
        except SystemExit:
            print("[Worker Thread] SystemExit - re-raising")
            raise  # Don't catch system exit
        except BaseException as e:
            # Catch-all for ANY exception including fatal ones
            error_msg = str(e)
            error_type = type(e).__name__
            traceback_str = traceback.format_exc()
            
            print(f"\n[Worker Thread] CRITICAL ERROR: {error_type}")
            print(f"[Worker Thread] Error message: {error_msg}")
            print(f"[Worker Thread] Full traceback:\n{traceback_str}")
            
            # Emit error signal with full details
            detailed_error = (
                f"Lỗi nghiêm trọng trong worker thread!\n\n"
                f"Loại lỗi: {error_type}\n"
                f"Thông báo: {error_msg}\n\n"
                f"Full traceback:\n{traceback_str}"
            )
            self.error.emit(detailed_error)
            
            # Also try to trigger exception hook if possible
            import sys
            if sys.excepthook != sys.__excepthook__:
                try:
                    # Try to call exception hook
                    sys.excepthook(type(e), e, e.__traceback__)
                except:
                    pass
        finally:
            print("[Worker Thread] Cleaning up...")
            self.progress.emit("")


class BatchWorker(QThread):
    """Worker thread for batch processing multiple files"""
    finished = pyqtSignal()  # Batch processing complete
    file_finished = pyqtSignal(str, str, str)  # file_path, file_type, ocr_result
    file_error = pyqtSignal(str, str)  # file_path, error_message
    progress = pyqtSignal(int, int, str)  # current, total, current_file_name
    
    def __init__(self, model, processor, file_list, prompt, generation_params):
        super().__init__()
        self.model = model
        self.processor = processor
        self.file_list = file_list  # List of (file_path, file_type, image_path)
        self.prompt = prompt
        self.generation_params = generation_params
        self.should_stop = False
    
    def stop(self):
        """Stop batch processing"""
        self.should_stop = True
    
    def run(self):
        """Process files sequentially"""
        total = len(self.file_list)
        
        for index, (file_path, file_type, image_path) in enumerate(self.file_list):
            if self.should_stop:
                break
            
            try:
                file_name = os.path.basename(file_path)
                self.progress.emit(index + 1, total, file_name)
                
                # Process image với model
                from PIL import Image
                
                # Load image
                if file_type == 'image':
                    image = Image.open(image_path).convert("RGB")
                else:
                    # PDF/DOCX/TXT đã được convert sang image
                    image = Image.open(image_path).convert("RGB")
                
                # Prepare messages
                messages = [
                    {
                        "role": "user",
                        "content": [
                            {"type": "image", "image": image},
                            {"type": "text", "text": self.prompt}
                        ]
                    }
                ]
                
                # Prepare inputs
                inputs = self.processor.apply_chat_template(
                    messages,
                    tokenize=True,
                    add_generation_prompt=True,
                    return_dict=True,
                    return_tensors="pt"
                )
                inputs = inputs.to(self.model.device)
                
                # Generate
                generated_ids = self.model.generate(**inputs, **self.generation_params)
                
                # Decode
                generated_ids_trimmed = [
                    out_ids[len(in_ids):]
                    for in_ids, out_ids in zip(inputs.input_ids, generated_ids)
                ]
                output_text = self.processor.batch_decode(
                    generated_ids_trimmed,
                    skip_special_tokens=True,
                    clean_up_tokenization_spaces=False
                )
                
                result = output_text[0] if output_text else ""
                
                # Emit file finished signal
                self.file_finished.emit(file_path, file_type, result)
                
            except Exception as e:
                error_msg = f"Lỗi khi xử lý {file_name}: {str(e)}"
                print(f"[Batch Worker] {error_msg}")
                self.file_error.emit(file_path, error_msg)
        
        # Emit finished signal
        self.finished.emit()


class ModelWorker(QThread):
    """Worker thread for running model inference"""
    finished = pyqtSignal(str)
    error = pyqtSignal(str)
    progress = pyqtSignal(str)
    
    def __init__(self, model, processor, image_path, prompt, generation_params):
        super().__init__()
        self.model = model
        self.processor = processor
        self.image_path = image_path
        self.prompt = prompt
        self.generation_params = generation_params
        
    def run(self):
        try:
            self.progress.emit("Đang xử lý hình ảnh...")
            
            # Prepare messages
            messages = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "image": self.image_path,
                        },
                        {"type": "text", "text": self.prompt},
                    ],
                }
            ]
            
            self.progress.emit("Đang chuẩn bị input cho model...")
            
            # Preparation for inference
            inputs = self.processor.apply_chat_template(
                messages,
                tokenize=True,
                add_generation_prompt=True,
                return_dict=True,
                return_tensors="pt"
            )
            inputs = inputs.to(self.model.device)
            
            self.progress.emit("Đang tạo output...")
            
            # Inference: Generation of the output
            generated_ids = self.model.generate(**inputs, **self.generation_params)
            generated_ids_trimmed = [
                out_ids[len(in_ids):] for in_ids, out_ids in zip(inputs.input_ids, generated_ids)
            ]
            output_text = self.processor.batch_decode(
                generated_ids_trimmed, skip_special_tokens=True, clean_up_tokenization_spaces=False
            )
            
            self.finished.emit(output_text[0] if output_text else "")
            
        except Exception as e:
            self.error.emit(f"Lỗi: {str(e)}")


class Qwen3VLApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.model = None
        self.processor = None
        self.current_image_path = None
        self.current_file_path = None
        self.current_file_type = None  # 'image', 'pdf', 'docx', 'txt'
        self.temp_image_files = []  # Track temp files for PDF/DOCX
        self.worker = None
        self.model_loader_worker = None  # For model loading
        self.current_device = "auto"  # Track current device selection
        self.current_model = "4B"  # Track current model selection (2B or 4B)
        self.local_model_path = "./models/Qwen3-VL-4B-Instruct"  # Default to 4B
        self.auto_load_enabled = False  # Tắt auto-load mặc định để tránh crash
        self.progress_dialog = None  # Progress dialog for model loading
        self.processing_start_time = None  # Track processing start time
        
        # SQLite database
        self.db_path = "./ocr_history.db"
        self.history_dir = "./history"
        self.init_database()
        
        # Batch processing
        self.batch_queue = []  # Queue các files cần xử lý
        self.batch_worker = None  # Worker thread cho batch processing
        self.current_batch_index = 0
        self.total_batch_files = 0
        self.batch_results = []  # Lưu kết quả từng file
        
        # ROI processing
        self.current_processing_image = None  # Image path đang xử lý (có thể là cropped)
        self.processing_temp_files = []  # Temp files từ ROI crop
        
        self.init_ui()
        
        # Không auto-load mặc định - user phải click nút "Sử dụng AI này" thủ công
        # Auto-load có thể gây crash nếu memory không đủ
        # if self.auto_load_enabled:
        #     from PyQt6.QtCore import QTimer
        #     QTimer.singleShot(2000, self.auto_load_model)
    
    def apply_modern_stylesheet(self):
        """Apply modern stylesheet to the application"""
        modern_stylesheet = """
        /* Main Window */
        QMainWindow {
            background-color: #f5f7fa;
        }
        
        /* Group Boxes */
        QGroupBox {
            font-size: 12pt;
            font-weight: bold;
            color: #2c3e50;
            border: 2px solid #e1e8ed;
            border-radius: 12px;
            margin-top: 12px;
            padding-top: 15px;
            background-color: white;
        }
        
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 15px;
            padding: 0 8px 0 8px;
            color: #3498db;
        }
        
        /* Labels */
        QLabel {
            color: #2c3e50;
            font-size: 10pt;
        }
        
        /* Đảm bảo tất cả text widget có màu đen */
        QLineEdit, QPlainTextEdit {
            color: #2c3e50;
            background-color: white;
        }
        
        /* Menu và context menu */
        QMenu {
            background-color: white;
            color: #2c3e50;
            border: 1px solid #e1e8ed;
            border-radius: 4px;
        }
        
        QMenu::item {
            color: #2c3e50;
            padding: 5px 20px;
        }
        
        QMenu::item:selected {
            background-color: #3498db;
            color: white;
        }
        
        /* Buttons */
        QPushButton {
            background-color: white;
            color: #2c3e50;
            border: 2px solid #3498db;
            border-radius: 8px;
            padding: 10px 20px;
            font-size: 11pt;
            font-weight: bold;
            min-height: 35px;
        }
        
        QPushButton:hover {
            background-color: #ecf0f1;
            border-color: #2980b9;
        }
        
        QPushButton:pressed {
            background-color: #d5dbdb;
            border-color: #21618c;
        }
        
        QPushButton:disabled {
            background-color: #ecf0f1;
            color: #95a5a6;
            border-color: #bdc3c7;
        }
        
        /* Process Button - Special Style */
        QPushButton[objectName="process_btn"] {
            background-color: white;
            color: #2c3e50;
            border: 2px solid #27ae60;
            font-size: 14pt;
            padding: 15px 30px;
            border-radius: 10px;
        }
        
        QPushButton[objectName="process_btn"]:hover {
            background-color: #ecf0f1;
            border-color: #229954;
        }
        
        QPushButton[objectName="process_btn"]:pressed {
            background-color: #d5dbdb;
            border-color: #1e8449;
        }
        
        QPushButton[objectName="process_btn"]:disabled {
            background-color: #ecf0f1;
            color: #95a5a6;
            border-color: #bdc3c7;
        }
        
        /* ComboBox */
        QComboBox {
            background-color: white;
            border: 2px solid #e1e8ed;
            border-radius: 8px;
            padding: 8px 15px;
            font-size: 10pt;
            min-height: 35px;
            color: #2c3e50;
        }
        
        QComboBox:hover {
            border: 2px solid #3498db;
        }
        
        QComboBox:focus {
            border: 2px solid #3498db;
            outline: none;
        }
        
        QComboBox::drop-down {
            border: none;
            width: 30px;
        }
        
        QComboBox::down-arrow {
            image: none;
            border-left: 5px solid transparent;
            border-right: 5px solid transparent;
            border-top: 6px solid #2c3e50;
            width: 0;
            height: 0;
            margin-right: 10px;
        }
        
        QComboBox QAbstractItemView {
            background-color: white;
            color: #2c3e50;
            border: 2px solid #e1e8ed;
            border-radius: 8px;
            selection-background-color: #3498db;
            selection-color: white;
        }
        
        /* SpinBox and DoubleSpinBox */
        QSpinBox, QDoubleSpinBox {
            background-color: white;
            border: 2px solid #e1e8ed;
            border-radius: 8px;
            padding: 8px 15px;
            font-size: 10pt;
            min-height: 35px;
            color: #2c3e50;
        }
        
        QSpinBox:hover, QDoubleSpinBox:hover {
            border: 2px solid #3498db;
        }
        
        QSpinBox:focus, QDoubleSpinBox:focus {
            border: 2px solid #3498db;
        }
        
        /* TextEdit */
        QTextEdit {
            background-color: white;
            border: 2px solid #e1e8ed;
            border-radius: 8px;
            padding: 10px;
            font-size: 10pt;
            color: #2c3e50;
            selection-background-color: #3498db;
            selection-color: white;
        }
        
        QTextEdit:focus {
            border: 2px solid #3498db;
            outline: none;
        }
        
        /* Progress Bar */
        QProgressBar {
            border: 2px solid #e1e8ed;
            border-radius: 8px;
            text-align: center;
            font-size: 10pt;
            font-weight: bold;
            background-color: #ecf0f1;
            height: 25px;
            color: #2c3e50;
        }
        
        QProgressBar::chunk {
            background-color: #3498db;
            border-radius: 6px;
        }
        
        /* CheckBox */
        QCheckBox {
            font-size: 10pt;
            color: #2c3e50;
            spacing: 8px;
        }
        
        QCheckBox::indicator {
            width: 20px;
            height: 20px;
            border: 2px solid #3498db;
            border-radius: 4px;
            background-color: white;
        }
        
        QCheckBox::indicator:hover {
            border: 2px solid #2980b9;
        }
        
        QCheckBox::indicator:checked {
            background-color: #3498db;
            border: 2px solid #3498db;
            image: none;
        }
        
        QCheckBox::indicator:checked {
            background-color: #3498db;
            border: 2px solid #3498db;
        }
        
        /* ScrollBar */
        QScrollBar:vertical {
            border: none;
            background-color: #ecf0f1;
            width: 12px;
            border-radius: 6px;
        }
        
        QScrollBar::handle:vertical {
            background-color: #95a5a6;
            border-radius: 6px;
            min-height: 30px;
        }
        
        QScrollBar::handle:vertical:hover {
            background-color: #7f8c8d;
        }
        
        QScrollBar:horizontal {
            border: none;
            background-color: #ecf0f1;
            height: 12px;
            border-radius: 6px;
        }
        
        QScrollBar::handle:horizontal {
            background-color: #95a5a6;
            border-radius: 6px;
            min-width: 30px;
        }
        
        QScrollBar::handle:horizontal:hover {
            background-color: #7f8c8d;
        }
        """
        
        self.setStyleSheet(modern_stylesheet)
    
    def init_database(self):
        """Khởi tạo SQLite database và tạo bảng lịch sử OCR"""
        try:
            # Tạo thư mục history nếu chưa có
            os.makedirs(self.history_dir, exist_ok=True)
            
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Tạo bảng lịch sử OCR
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS ocr_history (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_path TEXT NOT NULL,
                    file_name TEXT NOT NULL,
                    file_type TEXT NOT NULL,
                    ocr_result TEXT NOT NULL,
                    ocr_result_json TEXT,
                    preview_image_path TEXT,
                    timestamp TEXT NOT NULL,
                    model_used TEXT,
                    processing_time REAL
                )
            ''')
            
            conn.commit()
            conn.close()
            print(f"[Database] Đã khởi tạo database: {self.db_path}")
        except Exception as e:
            print(f"[Database] Lỗi khi khởi tạo database: {e}")
    
    def create_thumbnail(self, image_path, max_size=(200, 200)):
        """Tạo thumbnail/preview image"""
        try:
            if not os.path.exists(image_path):
                return None
            
            # Mở và resize image
            img = Image.open(image_path)
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
            
            # Lưu thumbnail vào history directory
            file_name = os.path.basename(image_path)
            name, ext = os.path.splitext(file_name)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            thumbnail_path = os.path.join(self.history_dir, f"thumb_{timestamp}_{name}{ext}")
            
            img.save(thumbnail_path, quality=85)
            return thumbnail_path
        except Exception as e:
            print(f"[Thumbnail] Lỗi khi tạo thumbnail: {e}")
            return None
    
    def save_history(self, file_path, file_type, ocr_result, processing_time=None):
        """Lưu kết quả OCR vào database"""
        try:
            file_name = os.path.basename(file_path)
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            model_used = f"{self.current_model}B" if self.current_model else "Unknown"
            
            # Tạo thumbnail cho preview
            preview_image_path = None
            if file_type == 'image':
                preview_image_path = self.create_thumbnail(file_path)
            elif file_type in ['pdf', 'docx', 'txt']:
                # Đối với PDF/DOCX/TXT, dùng ảnh đã được tạo sẵn nếu có
                if self.current_image_path and os.path.exists(self.current_image_path):
                    preview_image_path = self.create_thumbnail(self.current_image_path)
            
            # Try to parse JSON from OCR result
            ocr_result_json = None
            try:
                # Try to extract JSON from result (in case there's extra text)
                result_text = ocr_result.strip()
                # Find JSON object in the result
                if '{' in result_text and '}' in result_text:
                    json_start = result_text.find('{')
                    json_end = result_text.rfind('}') + 1
                    json_str = result_text[json_start:json_end]
                    # Validate it's valid JSON
                    import json
                    parsed = json.loads(json_str)
                    ocr_result_json = json.dumps(parsed, ensure_ascii=False, indent=2)
                    print(f"[Database] Đã phân tích JSON từ kết quả OCR")
            except Exception as json_error:
                print(f"[Database] Không thể phân tích JSON từ kết quả: {json_error}")
                # Keep ocr_result_json as None if parsing fails
            
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT INTO ocr_history 
                (file_path, file_name, file_type, ocr_result, ocr_result_json, preview_image_path, 
                 timestamp, model_used, processing_time)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (file_path, file_name, file_type, ocr_result, ocr_result_json, preview_image_path,
                  timestamp, model_used, processing_time))
            
            conn.commit()
            history_id = cursor.lastrowid
            conn.close()
            
            print(f"[Database] Đã lưu lịch sử OCR: ID={history_id}, File={file_name}")
            
            # Refresh history tab nếu đang mở
            self.refresh_history()
            
            return history_id
        except Exception as e:
            print(f"[Database] Lỗi khi lưu lịch sử: {e}")
            return None
    
    def load_history(self, limit=100):
        """Load lịch sử OCR từ database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT id, file_path, file_name, file_type, ocr_result, ocr_result_json,
                       preview_image_path, timestamp, model_used, processing_time
                FROM ocr_history
                ORDER BY timestamp DESC
                LIMIT ?
            ''', (limit,))
            
            results = cursor.fetchall()
            conn.close()
            
            history_list = []
            for row in results:
                history_list.append({
                    'id': row[0],
                    'file_path': row[1],
                    'file_name': row[2],
                    'file_type': row[3],
                    'ocr_result': row[4],
                    'ocr_result_json': row[5],
                    'preview_image_path': row[6],
                    'timestamp': row[7],
                    'model_used': row[8],
                    'processing_time': row[9]
                })
            
            return history_list
        except Exception as e:
            print(f"[Database] Lỗi khi load lịch sử: {e}")
            return []
    
    def update_history(self, history_id, ocr_result=None, file_name=None):
        """Update lịch sử OCR trong database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            updates = []
            values = []
            
            if ocr_result is not None:
                updates.append("ocr_result = ?")
                values.append(ocr_result)
            
            if file_name is not None:
                updates.append("file_name = ?")
                values.append(file_name)
            
            if not updates:
                conn.close()
                return False
            
            # Update timestamp khi sửa
            updates.append("timestamp = ?")
            values.append(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            
            values.append(history_id)
            
            cursor.execute(f'''
                UPDATE ocr_history
                SET {', '.join(updates)}
                WHERE id = ?
            ''', tuple(values))
            
            conn.commit()
            conn.close()
            
            print(f"[Database] Đã cập nhật lịch sử OCR: ID={history_id}")
            
            # Refresh history tab
            self.refresh_history()
            
            return True
        except Exception as e:
            print(f"[Database] Lỗi khi cập nhật lịch sử: {e}")
            return False
    
    def delete_history(self, history_id):
        """Xóa một lịch sử OCR từ database"""
        try:
            # Load thông tin trước khi xóa để xóa preview image
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('SELECT preview_image_path FROM ocr_history WHERE id = ?', (history_id,))
            result = cursor.fetchone()
            
            if result and result[0] and os.path.exists(result[0]):
                try:
                    os.remove(result[0])
                    print(f"[Database] Đã xóa preview image: {result[0]}")
                except:
                    pass
            
            # Xóa record
            cursor.execute('DELETE FROM ocr_history WHERE id = ?', (history_id,))
            conn.commit()
            conn.close()
            
            print(f"[Database] Đã xóa lịch sử OCR: ID={history_id}")
            
            # Refresh history tab
            self.refresh_history()
            
            return True
        except Exception as e:
            print(f"[Database] Lỗi khi xóa lịch sử: {e}")
            return False
    
    def delete_multiple_history(self, history_ids):
        """Xóa nhiều lịch sử OCR từ database"""
        try:
            if not history_ids:
                return False
            
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Load preview images trước khi xóa
            placeholders = ','.join(['?'] * len(history_ids))
            cursor.execute(f'SELECT preview_image_path FROM ocr_history WHERE id IN ({placeholders})', history_ids)
            results = cursor.fetchall()
            
            # Xóa preview images
            for row in results:
                if row[0] and os.path.exists(row[0]):
                    try:
                        os.remove(row[0])
                    except:
                        pass
            
            # Xóa records
            cursor.execute(f'DELETE FROM ocr_history WHERE id IN ({placeholders})', history_ids)
            conn.commit()
            conn.close()
            
            print(f"[Database] Đã xóa {len(history_ids)} lịch sử OCR")
            
            # Refresh history tab
            self.refresh_history()
            
            return True
        except Exception as e:
            print(f"[Database] Lỗi khi xóa nhiều lịch sử: {e}")
            return False
    
    def refresh_history(self):
        """Refresh danh sách lịch sử trong tab History"""
        if hasattr(self, 'history_list_widget') and self.history_list_widget:
            self.history_list_widget.clear()
            history = self.load_history()
            
            for item in history:
                # Tạo item text với timestamp và file name
                display_text = f"{item['timestamp']} - {item['file_name']}"
                if item['model_used']:
                    display_text += f" ({item['model_used']})"
                
                list_item = QListWidgetItem(display_text)
                list_item.setData(Qt.ItemDataRole.UserRole, item)  # Lưu full data
                
                # Nếu có preview image, set icon
                if item['preview_image_path'] and os.path.exists(item['preview_image_path']):
                    try:
                        pixmap = QPixmap(item['preview_image_path'])
                        if not pixmap.isNull():
                            icon = QPixmap(pixmap).scaled(64, 64, Qt.AspectRatioMode.KeepAspectRatio, 
                                                         Qt.TransformationMode.SmoothTransformation)
                            list_item.setIcon(icon)
                    except:
                        pass
                
                self.history_list_widget.addItem(list_item)
        
    def init_ui(self):
        self.setWindowTitle("Ứng Dụng OCR")
        
        # Fullscreen on startup - full màn hình máy tính
        self.showFullScreen()
        
        # Set minimum window size for responsive design - optimized for usability
        self.setMinimumSize(900, 600)
        
        # Apply modern stylesheet
        self.apply_modern_stylesheet()
        
        # Central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        central_widget.setStyleSheet("background-color: #f5f7fa;")
        
        # Main layout
        main_layout = QVBoxLayout()
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        central_widget.setLayout(main_layout)
        
        # Top control bar for fullscreen mode
        control_bar = QWidget()
        control_bar.setStyleSheet("background-color: #2c3e50; padding: 5px;")
        control_bar_layout = QHBoxLayout()
        control_bar_layout.setContentsMargins(10, 5, 10, 5)
        control_bar.setLayout(control_bar_layout)
        
        # App title
        app_title = QLabel("Ứng Dụng OCR")
        app_title.setStyleSheet("color: white; font-size: 14pt; font-weight: bold;")
        control_bar_layout.addWidget(app_title)
        
        control_bar_layout.addStretch()
        
        # Exit fullscreen button
        self.exit_fullscreen_btn = QPushButton("🗗 Thoát Fullscreen")
        self.exit_fullscreen_btn.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 8px 15px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        self.exit_fullscreen_btn.clicked.connect(self.exit_fullscreen)
        self.exit_fullscreen_btn.setToolTip("Nhấn để thoát chế độ fullscreen (hoặc nhấn ESC)")
        control_bar_layout.addWidget(self.exit_fullscreen_btn)
        
        # Close app button
        close_app_btn = QPushButton("✕ Đóng Ứng Dụng")
        close_app_btn.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 8px 15px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        close_app_btn.clicked.connect(self.close)
        close_app_btn.setToolTip("Đóng ứng dụng (hoặc nhấn Alt+F4)")
        control_bar_layout.addWidget(close_app_btn)
        
        main_layout.addWidget(control_bar)
        
        # Tab widget
        self.tab_widget = QTabWidget()
        self.tab_widget.setStyleSheet("""
            QTabWidget::pane {
                border: 1px solid #e1e8ed;
                background-color: #f5f7fa;
            }
            QTabBar::tab {
                background-color: white;
                color: #2c3e50;
                padding: 10px 20px;
                border: 1px solid #e1e8ed;
                border-bottom: none;
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
            }
            QTabBar::tab:selected {
                background-color: #3498db;
                color: white;
                border-color: #3498db;
            }
            QTabBar::tab:hover {
                background-color: #ecf0f1;
            }
        """)
        main_layout.addWidget(self.tab_widget)
        
        # Tab 1: OCR
        ocr_tab = QWidget()
        ocr_tab.setStyleSheet("background-color: #f5f7fa;")
        
        # Use QSplitter for resizable panels
        ocr_splitter = QSplitter(Qt.Orientation.Horizontal)
        ocr_splitter.setHandleWidth(4)
        ocr_splitter.setChildrenCollapsible(False)  # Prevent panels from collapsing
        ocr_splitter.setStyleSheet("""
            QSplitter::handle {
                background-color: #e1e8ed;
                width: 4px;
                margin: 0px;
                padding: 0px;
            }
            QSplitter::handle:hover {
                background-color: #3498db;
            }
        """)
        
        # Column 1 (Left) - Model configuration only with scroll
        left_scroll = QScrollArea()
        left_scroll.setWidgetResizable(True)
        left_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        left_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        left_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        left_scroll.setStyleSheet("QScrollArea { border: none; background-color: #f5f7fa; }")
        
        left_panel_widget = QWidget()
        left_panel_widget.setStyleSheet("background-color: #f5f7fa;")
        left_panel_widget.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.MinimumExpanding)
        left_panel = QVBoxLayout()
        left_panel.setSpacing(15)
        left_panel.setContentsMargins(10, 10, 10, 10)
        
        # Model loading section
        model_group = QGroupBox("Cấu Hình Model")
        model_layout = QVBoxLayout()
        
        # Model selection
        model_selection_layout = QHBoxLayout()
        model_selection_layout.addWidget(QLabel("Model:"))
        self.model_combo = QComboBox()
        self.model_combo.addItem("Model OCR mặc định", "4B")
        self.model_combo.addItem("Model OCR nhẹ", "2B")
        self.model_combo.setCurrentIndex(0)  # Default to 4B
        self.model_combo.currentIndexChanged.connect(self.on_model_changed)
        self.model_combo.setToolTip("Chọn model: OCR mặc định (chính xác hơn) hoặc OCR nhẹ (nhanh hơn)")
        model_selection_layout.addWidget(self.model_combo)
        model_layout.addLayout(model_selection_layout)
        
        self.model_label = QLabel("Model: Chưa tải")
        self.model_label.setStyleSheet("color: #2c3e50; font-size: 9pt; font-weight: bold; padding: 5px;")
        model_layout.addWidget(self.model_label)
        
        # Update model label based on initial selection
        self.update_model_label()
        
        # Auto-load checkbox
        auto_load_layout = QHBoxLayout()
        self.auto_load_checkbox = QCheckBox("Tự động load model khi khởi động (Khuyến nghị: TẮT)")
        self.auto_load_checkbox.setChecked(False)  # Tắt mặc định để tránh crash
        self.auto_load_checkbox.stateChanged.connect(self.on_auto_load_changed)
        auto_load_layout.addWidget(self.auto_load_checkbox)
        model_layout.addLayout(auto_load_layout)
        
        # Device selection
        device_layout = QHBoxLayout()
        device_layout.addWidget(QLabel("Thiết bị:"))
        self.device_combo = QComboBox()
        
        # Check GPU availability - lazy check để tăng tốc khởi động
        cuda_info = get_cuda_info()
        if cuda_info['available']:
            gpu_name = cuda_info['device_name']
            self.device_combo.addItem(f"GPU ({gpu_name})", "cuda")
            self.device_combo.addItem("CPU", "cpu")
            self.device_combo.setCurrentIndex(0)  # Default to GPU if available
            self.device_combo.setToolTip("Khuyến nghị: Chọn GPU để có kết quả nhanh và tốt hơn CPU")
        else:
            self.device_combo.addItem("CPU", "cpu")
            self.device_combo.setCurrentIndex(0)
            self.device_combo.setToolTip("Không có GPU - Sẽ sử dụng CPU (chậm hơn GPU)")
        
        self.device_combo.currentIndexChanged.connect(self.on_device_changed)
        device_layout.addWidget(self.device_combo)
        model_layout.addLayout(device_layout)
        
        # Device info label - lazy check
        self.device_info = QLabel("")
        cuda_info = get_cuda_info()
        if cuda_info['available']:
            self.device_info.setText(f"GPU: {cuda_info['device_name']} | CUDA: {cuda_info['cuda_version']}")
        else:
            self.device_info.setText("Chế độ CPU - Không có GPU")
        self.device_info.setStyleSheet("color: blue; font-size: 9pt;")
        model_layout.addWidget(self.device_info)
        
        # Device recommendation label - lazy check
        self.device_recommendation = QLabel("")
        if cuda_info['available']:
            self.device_recommendation.setText("Khuyến nghị: Sử dụng GPU để có kết quả nhanh và tốt hơn CPU")
            self.device_recommendation.setStyleSheet("color: green; font-size: 9pt; font-weight: bold; padding: 5px;")
        else:
            self.device_recommendation.setText("Không có GPU - Sẽ sử dụng CPU (chậm hơn)")
            self.device_recommendation.setStyleSheet("color: orange; font-size: 9pt; font-weight: bold; padding: 5px;")
        model_layout.addWidget(self.device_recommendation)
        
        # Buttons layout
        buttons_layout = QHBoxLayout()
        buttons_layout.setSpacing(10)
        self.load_model_btn = QPushButton("Sử dụng AI này")
        self.load_model_btn.clicked.connect(self.load_model)
        buttons_layout.addWidget(self.load_model_btn)
        
        self.unload_model_btn = QPushButton("Gỡ Model")
        self.unload_model_btn.clicked.connect(self.unload_model)
        self.unload_model_btn.setEnabled(False)
        buttons_layout.addWidget(self.unload_model_btn)
        model_layout.addLayout(buttons_layout)
        
        self.model_status = QLabel("Trạng thái: Chưa tải model")
        self.model_status.setStyleSheet("""
            color: #e74c3c;
            font-size: 10pt;
            font-weight: bold;
            padding: 8px;
            background-color: white;
            border: 2px solid #e1e8ed;
            border-radius: 8px;
        """)
        model_layout.addWidget(self.model_status)
        
        model_group.setLayout(model_layout)
        left_panel.addWidget(model_group)
        
        # Add stretch to push content to top
        left_panel.addStretch()
        left_panel_widget.setLayout(left_panel)
        
        # Set scroll area widget
        left_scroll.setWidget(left_panel_widget)
        
        # Set minimum width for left panel
        left_scroll.setMinimumWidth(250)
        left_scroll.setMaximumWidth(400)
        left_scroll.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Expanding)
        
        # Column 2 (Middle) - File preview, ROI, and upload controls with scroll
        middle_scroll = QScrollArea()
        middle_scroll.setWidgetResizable(True)
        middle_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        middle_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        middle_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        middle_scroll.setStyleSheet("QScrollArea { border: none; background-color: #f5f7fa; }")
        
        middle_panel_widget = QWidget()
        middle_panel_widget.setStyleSheet("background-color: #f5f7fa;")
        middle_panel_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.MinimumExpanding)
        middle_panel = QVBoxLayout()
        middle_panel.setSpacing(15)
        middle_panel.setContentsMargins(10, 10, 10, 10)
        
        # File display
        file_group = QGroupBox("Xem Trước File")
        file_layout = QVBoxLayout()
        
        # ROI mode controls
        roi_controls_layout = QHBoxLayout()
        roi_controls_layout.addWidget(QLabel("Chế độ OCR:"))
        
        self.roi_mode_checkbox = QCheckBox("Chọn vùng tùy chỉnh (ROI)")
        self.roi_mode_checkbox.setChecked(False)
        self.roi_mode_checkbox.stateChanged.connect(self.on_roi_mode_changed)
        roi_controls_layout.addWidget(self.roi_mode_checkbox)
        
        self.clear_roi_btn = QPushButton("Xóa Vùng Chọn")
        self.clear_roi_btn.clicked.connect(self.clear_roi_selection)
        self.clear_roi_btn.setEnabled(False)
        roi_controls_layout.addWidget(self.clear_roi_btn)
        
        file_layout.addLayout(roi_controls_layout)
        
        self.image_label = ImageROILabel()
        self.image_label.setText("Chưa có file nào được tải")
        self.image_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        # Make image label responsive - very flexible sizing
        self.image_label.setMinimumSize(200, 200)
        self.image_label.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.image_label.setScaledContents(False)  # Maintain aspect ratio
        self.image_label.setStyleSheet("""
            border: 3px solid #e1e8ed;
            border-radius: 12px;
            background-color: white;
            padding: 10px;
        """)
        self.image_label.roi_changed.connect(self.on_roi_changed)
        file_layout.addWidget(self.image_label)
        
        # ROI info label
        self.roi_info_label = QLabel("Chưa chọn vùng. Click và drag để chọn vùng OCR")
        self.roi_info_label.setStyleSheet("color: #3498db; font-size: 9pt; padding: 5px;")
        self.roi_info_label.setVisible(False)  # Hidden by default, shown when ROI mode is enabled
        file_layout.addWidget(self.roi_info_label)
        
        # File info label
        self.file_info_label = QLabel("File: Không có")
        self.file_info_label.setStyleSheet("color: #7f8c8d; font-size: 9pt; font-weight: bold; padding: 5px;")
        file_layout.addWidget(self.file_info_label)
        
        # File buttons
        file_buttons_layout = QHBoxLayout()
        
        self.load_file_btn = QPushButton("Tải File")
        self.load_file_btn.clicked.connect(self.load_file)
        file_buttons_layout.addWidget(self.load_file_btn)
        
        self.load_multiple_files_btn = QPushButton("Chọn Nhiều File")
        self.load_multiple_files_btn.clicked.connect(self.load_multiple_files)
        file_buttons_layout.addWidget(self.load_multiple_files_btn)
        
        file_layout.addLayout(file_buttons_layout)
        
        # Batch processing info
        self.batch_info_label = QLabel("")
        self.batch_info_label.setStyleSheet("color: #2c3e50; font-size: 9pt; padding: 5px;")
        file_layout.addWidget(self.batch_info_label)
        
        # Batch progress bar
        self.batch_progress_bar = QProgressBar()
        self.batch_progress_bar.setVisible(False)
        file_layout.addWidget(self.batch_progress_bar)
        
        # File type label
        self.file_type_label = QLabel("Hỗ trợ: JPG, PNG, JPEG, PDF, DOCX, TXT")
        self.file_type_label.setStyleSheet("color: #3498db; font-size: 8pt; font-weight: bold; padding: 3px;")
        file_layout.addWidget(self.file_type_label)
        
        file_group.setLayout(file_layout)
        middle_panel.addWidget(file_group)
        
        # Add stretch to push content to top
        middle_panel.addStretch()
        middle_panel_widget.setLayout(middle_panel)
        
        # Set scroll area widget
        middle_scroll.setWidget(middle_panel_widget)
        
        # Set minimum width for middle panel - more flexible
        middle_scroll.setMinimumWidth(280)
        middle_scroll.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        
        # Column 3 (Right) - Prompt and results with scroll
        right_scroll = QScrollArea()
        right_scroll.setWidgetResizable(True)
        right_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        right_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        right_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        right_scroll.setStyleSheet("QScrollArea { border: none; background-color: #f5f7fa; }")
        
        right_panel_widget = QWidget()
        right_panel_widget.setStyleSheet("background-color: #f5f7fa;")
        right_panel_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.MinimumExpanding)
        right_panel = QVBoxLayout()
        right_panel.setSpacing(15)
        right_panel.setContentsMargins(10, 10, 10, 10)
        
        # Generation parameters
        params_group = QGroupBox("Tham Số Generation")
        params_layout = QVBoxLayout()
        
        # Prompt selection
        prompt_layout = QHBoxLayout()
        prompt_layout.addWidget(QLabel("Tác vụ:"))
        self.prompt_combo = QComboBox()
        # New structured prompt for Vietnamese official documents
        default_prompt = '''You are an expert in analyzing Vietnamese official documents. Your task is to carefully read the document from the provided image(s) and extract the following specific information.

Please return the output in a structured JSON format. The JSON object should contain the following keys:

- "loai_van_ban": The type of the document (e.g., "Nghị định", "Quyết định", "Thông tư", "Công văn").
- "so_ky_hieu": The official reference number of the document.
- "trich_yeu": A brief summary or subject of the document's content ("Trích yếu nội dung").
- "co_quan_ban_hanh": The name of the issuing authority or organization.
- "ngay_ban_hanh": The date of issue in "dd/mm/yyyy" format.

If any information is not found, please return a null or empty string for that key. Do not add any extra text or explanations outside of the JSON object.'''
        self.prompt_combo.addItem(default_prompt)
        self.prompt_combo.setEditable(False)  # Không cho phép chỉnh sửa
        self.prompt_combo.currentTextChanged.connect(self.on_prompt_changed)
        
        # Nút điền prompt mặc định vào custom prompt
        use_default_btn = QPushButton("Sử dụng prompt mặc định")
        use_default_btn.clicked.connect(self.on_use_default_prompt)
        prompt_layout.addWidget(self.prompt_combo)
        prompt_layout.addWidget(use_default_btn)
        params_layout.addLayout(prompt_layout)
        
        # Custom prompt input
        self.custom_prompt = QTextEdit()
        self.custom_prompt.setPlaceholderText("Nhập prompt tùy chỉnh tại đây... (Để trống sẽ dùng prompt mặc định)")
        self.custom_prompt.setMaximumHeight(60)
        params_layout.addWidget(self.custom_prompt)
        
        # Max new tokens
        max_tokens_layout = QHBoxLayout()
        max_tokens_layout.addWidget(QLabel("Số tokens tối đa:"))
        self.max_tokens_spin = QSpinBox()
        self.max_tokens_spin.setRange(1, 16384)
        self.max_tokens_spin.setValue(3000)
        max_tokens_layout.addWidget(self.max_tokens_spin)
        params_layout.addLayout(max_tokens_layout)
        
        # Temperature
        temp_layout = QHBoxLayout()
        temp_layout.addWidget(QLabel("Temperature:"))
        self.temperature_spin = QDoubleSpinBox()
        self.temperature_spin.setRange(0.0, 2.0)
        self.temperature_spin.setSingleStep(0.1)
        self.temperature_spin.setValue(0.2)
        temp_layout.addWidget(self.temperature_spin)
        params_layout.addLayout(temp_layout)
        
        # Top-p
        top_p_layout = QHBoxLayout()
        top_p_layout.addWidget(QLabel("Top-p:"))
        self.top_p_spin = QDoubleSpinBox()
        self.top_p_spin.setRange(0.0, 1.0)
        self.top_p_spin.setSingleStep(0.1)
        self.top_p_spin.setValue(0.8)
        top_p_layout.addWidget(self.top_p_spin)
        params_layout.addLayout(top_p_layout)
        
        # Top-k
        top_k_layout = QHBoxLayout()
        top_k_layout.addWidget(QLabel("Top-k:"))
        self.top_k_spin = QSpinBox()
        self.top_k_spin.setRange(1, 100)
        self.top_k_spin.setValue(20)
        top_k_layout.addWidget(self.top_k_spin)
        params_layout.addLayout(top_k_layout)
        
        # Repetition penalty
        rep_penalty_layout = QHBoxLayout()
        rep_penalty_layout.addWidget(QLabel("Hệ số lặp lại:"))
        self.rep_penalty_spin = QDoubleSpinBox()
        self.rep_penalty_spin.setRange(1.0, 2.0)
        self.rep_penalty_spin.setSingleStep(0.1)
        self.rep_penalty_spin.setValue(1.0)
        rep_penalty_layout.addWidget(self.rep_penalty_spin)
        params_layout.addLayout(rep_penalty_layout)
        
        params_group.setLayout(params_layout)
        right_panel.addWidget(params_group, 1)  # Stretch factor 1 for params
        
        # Process button
        self.process_btn = QPushButton("Xử Lý Hình Ảnh")
        self.process_btn.setObjectName("process_btn")
        self.process_btn.clicked.connect(self.process_image)
        self.process_btn.setEnabled(False)
        right_panel.addWidget(self.process_btn)
        
        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        right_panel.addWidget(self.progress_bar)
        
        # Progress label
        self.progress_label = QLabel("")
        self.progress_label.setStyleSheet("color: #2c3e50; font-size: 10pt; padding: 5px;")
        right_panel.addWidget(self.progress_label)
        
        # Processing time label
        self.processing_time_label = QLabel("Thời gian xử lý: Chưa bắt đầu")
        self.processing_time_label.setStyleSheet("color: #7f8c8d; font-size: 11pt; font-weight: bold; padding: 8px; background-color: white; border-radius: 8px;")
        right_panel.addWidget(self.processing_time_label)
        
        # Output section - Make it bigger
        output_group = QGroupBox("Kết Quả")
        output_layout = QVBoxLayout()
        # Give more space to output section
        output_layout.setSpacing(10)
        
        self.output_text = QTextEdit()
        self.output_text.setReadOnly(True)
        self.output_text.setPlaceholderText("Kết quả sẽ hiển thị tại đây...")
        # Make output text expandable and set flexible minimum height
        self.output_text.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.output_text.setMinimumHeight(250)  # Very flexible minimum height
        # Enable word wrap for better text display
        self.output_text.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
        output_layout.addWidget(self.output_text)
        
        # Save output button
        self.save_output_btn = QPushButton("Lưu Kết Quả Ra File")
        self.save_output_btn.clicked.connect(self.save_output)
        output_layout.addWidget(self.save_output_btn)
        
        output_group.setLayout(output_layout)
        # Set stretch factor so output section takes more space
        right_panel.addWidget(output_group, 3)  # Stretch factor 3 for output
        
        # Add stretch to push content to top
        right_panel.addStretch(1)  # Smaller stretch for remaining space
        right_panel_widget.setLayout(right_panel)
        
        # Set scroll area widget
        right_scroll.setWidget(right_panel_widget)
        
        # Set minimum width for right panel - flexible sizing
        right_scroll.setMinimumWidth(320)
        right_scroll.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        
        # Add scroll areas to splitter (3 columns: left, middle, right)
        ocr_splitter.addWidget(left_scroll)
        ocr_splitter.addWidget(middle_scroll)
        ocr_splitter.addWidget(right_scroll)
        
        # Set stretch factors (left: 1, middle: 1.5, right: 2.5) - prioritize column 3
        ocr_splitter.setStretchFactor(0, 1)
        ocr_splitter.setStretchFactor(1, 2)
        ocr_splitter.setStretchFactor(2, 3)
        
        # Set initial sizes proportionally - balanced with more space for column 3
        # Column 1: 22%, Column 2: 32%, Column 3: 46%
        total_width = 1200  # Approximate default width
        ocr_splitter.setSizes([int(total_width * 0.22), int(total_width * 0.32), int(total_width * 0.46)])
        
        # Add splitter to OCR tab layout
        ocr_layout = QHBoxLayout()
        ocr_layout.setContentsMargins(0, 0, 0, 0)
        ocr_layout.setSpacing(0)
        ocr_layout.addWidget(ocr_splitter)
        ocr_tab.setLayout(ocr_layout)
        
        # Thêm OCR tab vào tab widget
        self.tab_widget.addTab(ocr_tab, "OCR")
        
        # Tab 2: History
        history_tab = QWidget()
        history_tab.setStyleSheet("background-color: #f5f7fa;")
        
        # Use QSplitter for resizable panels in History tab
        history_splitter = QSplitter(Qt.Orientation.Horizontal)
        history_splitter.setStyleSheet("""
            QSplitter::handle {
                background-color: #e1e8ed;
                width: 3px;
            }
            QSplitter::handle:hover {
                background-color: #3498db;
            }
        """)
        
        # History left panel widget - List
        history_left_widget = QWidget()
        history_left_widget.setStyleSheet("background-color: #f5f7fa;")
        history_left = QVBoxLayout()
        history_left.setContentsMargins(10, 10, 10, 10)
        history_left.setSpacing(15)
        history_left_group = QGroupBox("Lịch Sử OCR")
        history_left_layout = QVBoxLayout()
        
        self.history_list_widget = QListWidget()
        self.history_list_widget.setIconSize(QSize(64, 64))
        self.history_list_widget.itemClicked.connect(self.on_history_item_clicked)
        self.history_list_widget.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.history_list_widget.customContextMenuRequested.connect(self.on_history_context_menu)
        # Make history list responsive
        self.history_list_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        history_left_layout.addWidget(self.history_list_widget)
        
        # Buttons layout
        history_buttons_layout = QHBoxLayout()
        
        # Refresh button
        refresh_btn = QPushButton("Làm Mới")
        refresh_btn.clicked.connect(self.refresh_history)
        history_buttons_layout.addWidget(refresh_btn)
        
        # Delete selected button
        self.delete_history_btn = QPushButton("Xóa")
        self.delete_history_btn.clicked.connect(self.on_delete_history_clicked)
        self.delete_history_btn.setEnabled(False)
        history_buttons_layout.addWidget(self.delete_history_btn)
        
        # Delete all button
        delete_all_btn = QPushButton("Xóa Tất Cả")
        delete_all_btn.clicked.connect(self.on_delete_all_history)
        history_buttons_layout.addWidget(delete_all_btn)
        
        history_left_layout.addLayout(history_buttons_layout)
        
        history_left_group.setLayout(history_left_layout)
        history_left.addWidget(history_left_group)
        
        # Add stretch to push content to top
        history_left.addStretch()
        history_left_widget.setLayout(history_left)
        
        # Set minimum width for history left panel
        history_left_widget.setMinimumWidth(300)
        history_left_widget.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Preferred)
        
        # History right panel widget - Preview and Result
        history_right_widget = QWidget()
        history_right_widget.setStyleSheet("background-color: #f5f7fa;")
        history_right = QVBoxLayout()
        history_right.setContentsMargins(10, 10, 10, 10)
        history_right.setSpacing(15)
        history_right_group = QGroupBox("Chi Tiết")
        history_right_layout = QVBoxLayout()
        
        # Preview image
        self.history_preview_label = QLabel()
        # Make preview image responsive - flexible sizing
        self.history_preview_label.setMinimumSize(250, 250)
        self.history_preview_label.setMaximumSize(600, 600)
        self.history_preview_label.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        self.history_preview_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.history_preview_label.setStyleSheet("""
            border: 2px solid #e1e8ed;
            border-radius: 8px;
            background-color: white;
        """)
        self.history_preview_label.setText("Chọn một mục để xem preview")
        history_right_layout.addWidget(self.history_preview_label)
        
        # File info
        self.history_info_label = QLabel()
        self.history_info_label.setStyleSheet("color: #2c3e50; font-size: 10pt; padding: 5px;")
        history_right_layout.addWidget(self.history_info_label)
        
        # OCR result
        self.history_result_text = QTextEdit()
        self.history_result_text.setReadOnly(False)  # Cho phép edit
        self.history_result_text.setPlaceholderText("Chọn một mục để xem kết quả OCR...")
        self.history_result_text.textChanged.connect(self.on_history_result_changed)
        # Make history result text responsive
        self.history_result_text.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        history_right_layout.addWidget(self.history_result_text)
        
        # Edit buttons
        edit_buttons_layout = QHBoxLayout()
        
        # Save button
        self.save_history_btn = QPushButton("Lưu Thay Đổi")
        self.save_history_btn.clicked.connect(self.on_save_history_clicked)
        self.save_history_btn.setEnabled(False)
        edit_buttons_layout.addWidget(self.save_history_btn)
        
        # Cancel button
        self.cancel_edit_btn = QPushButton("Hủy")
        self.cancel_edit_btn.clicked.connect(self.on_cancel_edit_clicked)
        self.cancel_edit_btn.setEnabled(False)
        edit_buttons_layout.addWidget(self.cancel_edit_btn)
        
        history_right_layout.addLayout(edit_buttons_layout)
        
        # Track current editing history
        self.current_editing_history_id = None
        self.original_history_result = None
        
        history_right_group.setLayout(history_right_layout)
        history_right.addWidget(history_right_group)
        
        # Add stretch to push content to top
        history_right.addStretch()
        history_right_widget.setLayout(history_right)
        
        # Set minimum width for history right panel
        history_right_widget.setMinimumWidth(400)
        history_right_widget.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Preferred)
        
        # Add panels to splitter
        history_splitter.addWidget(history_left_widget)
        history_splitter.addWidget(history_right_widget)
        
        # Set stretch factors (left: 1, right: 2)
        history_splitter.setStretchFactor(0, 1)
        history_splitter.setStretchFactor(1, 2)
        
        # Set initial sizes (40% left, 60% right)
        history_splitter.setSizes([300, 500])
        
        # Add splitter to History tab layout
        history_layout = QHBoxLayout()
        history_layout.setContentsMargins(0, 0, 0, 0)
        history_layout.setSpacing(0)
        history_layout.addWidget(history_splitter)
        history_tab.setLayout(history_layout)
        
        # Thêm History tab vào tab widget
        self.tab_widget.addTab(history_tab, "Lịch Sử")
        
        # Load history lần đầu
        self.refresh_history()
        
    def on_history_item_clicked(self, item):
        """Xử lý khi click vào item trong history list"""
        data = item.data(Qt.ItemDataRole.UserRole)
        if not data:
            return
        
        # Hiển thị preview image
        if data.get('preview_image_path') and os.path.exists(data['preview_image_path']):
            try:
                pixmap = QPixmap(data['preview_image_path'])
                if not pixmap.isNull():
                    scaled_pixmap = pixmap.scaled(
                        self.history_preview_label.size(),
                        Qt.AspectRatioMode.KeepAspectRatio,
                        Qt.TransformationMode.SmoothTransformation
                    )
                    self.history_preview_label.setPixmap(scaled_pixmap)
                else:
                    self.history_preview_label.setText("Không thể load preview")
            except Exception as e:
                self.history_preview_label.setText(f"Lỗi load preview: {e}")
        else:
            self.history_preview_label.setText("Không có preview")
        
        # Hiển thị file info
        info_text = f"File: {data['file_name']}\n"
        info_text += f"Loại: {data['file_type'].upper()}\n"
        info_text += f"Thời gian: {data['timestamp']}\n"
        if data.get('model_used'):
            info_text += f"Model: {data['model_used']}\n"
        if data.get('processing_time'):
            info_text += f"Thời gian xử lý: {data['processing_time']:.2f}s"
        self.history_info_label.setText(info_text)
        
        # Hiển thị OCR result - ưu tiên JSON nếu có
        self.current_editing_history_id = data.get('id')
        
        # Try to display JSON if available
        if data.get('ocr_result_json'):
            self.original_history_result = data.get('ocr_result_json', '')
            self.history_result_text.setText(self.original_history_result)
        else:
            # Fallback to text result
            self.original_history_result = data.get('ocr_result', '')
            self.history_result_text.setText(self.original_history_result)
        
        # Enable/disable buttons
        self.save_history_btn.setEnabled(False)
        self.cancel_edit_btn.setEnabled(False)
        self.delete_history_btn.setEnabled(True)
    
    def on_history_result_changed(self):
        """Xử lý khi OCR result thay đổi"""
        if self.current_editing_history_id:
            current_text = self.history_result_text.toPlainText()
            has_changes = current_text != self.original_history_result
            self.save_history_btn.setEnabled(has_changes)
            self.cancel_edit_btn.setEnabled(has_changes)
    
    def on_save_history_clicked(self):
        """Lưu thay đổi OCR result"""
        if not self.current_editing_history_id:
            return
        
        new_result = self.history_result_text.toPlainText()
        
        if self.update_history(self.current_editing_history_id, ocr_result=new_result):
            QMessageBox.information(self, "Thành Công", "Đã lưu thay đổi thành công!")
            self.original_history_result = new_result
            self.save_history_btn.setEnabled(False)
            self.cancel_edit_btn.setEnabled(False)
        else:
            QMessageBox.warning(self, "Lỗi", "Không thể lưu thay đổi!")
    
    def on_cancel_edit_clicked(self):
        """Hủy thay đổi và khôi phục về giá trị ban đầu"""
        if self.current_editing_history_id and self.original_history_result is not None:
            self.history_result_text.setText(self.original_history_result)
            self.save_history_btn.setEnabled(False)
            self.cancel_edit_btn.setEnabled(False)
    
    def on_delete_history_clicked(self):
        """Xóa history item được chọn"""
        current_item = self.history_list_widget.currentItem()
        if not current_item:
            QMessageBox.warning(self, "Cảnh Báo", "Vui lòng chọn một mục để xóa!")
            return
        
        data = current_item.data(Qt.ItemDataRole.UserRole)
        if not data:
            return
        
        history_id = data.get('id')
        file_name = data.get('file_name', 'Unknown')
        
        reply = QMessageBox.question(
            self,
            "Xác Nhận Xóa",
            f"Bạn có chắc chắn muốn xóa '{file_name}'?\n\nThao tác này không thể hoàn tác!",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            if self.delete_history(history_id):
                QMessageBox.information(self, "Thành Công", "Đã xóa thành công!")
                # Clear selection
                self.history_preview_label.setText("Chọn một mục để xem preview")
                self.history_info_label.setText("")
                self.history_result_text.setText("")
                self.current_editing_history_id = None
                self.original_history_result = None
                self.delete_history_btn.setEnabled(False)
            else:
                QMessageBox.warning(self, "Lỗi", "Không thể xóa!")
    
    def on_delete_all_history(self):
        """Xóa tất cả history"""
        reply = QMessageBox.question(
            self,
            "Xác Nhận Xóa Tất Cả",
            "Bạn có chắc chắn muốn xóa TẤT CẢ lịch sử OCR?\n\nThao tác này không thể hoàn tác!",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            try:
                history = self.load_history(limit=1000)  # Load tất cả
                history_ids = [item['id'] for item in history]
                
                if history_ids:
                    if self.delete_multiple_history(history_ids):
                        QMessageBox.information(self, "Thành Công", f"Đã xóa {len(history_ids)} mục!")
                    else:
                        QMessageBox.warning(self, "Lỗi", "Không thể xóa tất cả!")
                else:
                    QMessageBox.information(self, "Thông Báo", "Không có dữ liệu để xóa!")
            except Exception as e:
                QMessageBox.warning(self, "Lỗi", f"Lỗi khi xóa: {e}")
    
    def on_history_context_menu(self, position):
        """Hiển thị context menu khi right-click vào history item"""
        item = self.history_list_widget.itemAt(position)
        if not item:
            return
        
        data = item.data(Qt.ItemDataRole.UserRole)
        if not data:
            return
        
        from PyQt6.QtWidgets import QMenu
        
        menu = QMenu(self)
        
        edit_action = menu.addAction("Chỉnh Sửa Kết Quả")
        delete_action = menu.addAction("Xóa")
        menu.addSeparator()
        view_action = menu.addAction("Xem Chi Tiết")
        
        action = menu.exec(self.history_list_widget.mapToGlobal(position))
        
        if action == edit_action:
            # Select item để có thể edit
            self.history_list_widget.setCurrentItem(item)
            self.on_history_item_clicked(item)
            # Focus vào text edit
            self.history_result_text.setFocus()
        
        elif action == delete_action:
            self.on_delete_history_clicked()
        
        elif action == view_action:
            self.on_history_item_clicked(item)
        
    def on_model_changed(self, index):
        """Handle model selection change"""
        selected_model = self.model_combo.currentData()
        
        # Warn if selecting 4B model on CPU
        if selected_model == "4B":
            current_device = self.device_combo.currentData()
            if current_device == "cpu":
                from PyQt6.QtWidgets import QMessageBox
                msg = QMessageBox(self)
                msg.setIcon(QMessageBox.Icon.Warning)
                msg.setWindowTitle("Cảnh Báo Model Mạnh Trên CPU")
                msg.setText("Model mạnh (4B) trên CPU sẽ RẤT CHẬM và NẶNG!")
                msg.setInformativeText(
                    "Model OCR mạnh (4B) không được khuyến nghị trên CPU vì:\n\n"
                    "• Rất chậm (có thể mất vài phút mỗi ảnh)\n"
                    "• Tiêu tốn RAM lớn (~10-12GB)\n"
                    "• Có thể gây treo máy\n\n"
                    "Khuyến nghị:\n"
                    "• Chuyển sang GPU (nếu có)\n"
                    "• Hoặc dùng Model nhẹ (2B) cho CPU\n\n"
                    "Bạn vẫn muốn dùng Model mạnh trên CPU?"
                )
                msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                msg.setDefaultButton(QMessageBox.StandardButton.No)
                
                if msg.exec() == QMessageBox.StandardButton.No:
                    # Revert to 2B model
                    self.model_combo.setCurrentIndex(1)  # Index 1 = 2B
                    return
        
        if self.model is not None:
            # Model đã được load, cần reload với model mới
            from PyQt6.QtWidgets import QMessageBox
            reply = QMessageBox.question(
                self,
                "Thay Đổi Model",
                f"Model đã được tải. Bạn có muốn chuyển sang model {selected_model} không?\n\n"
                "Lưu ý: Chuyển model sẽ cần tải lại model.",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                self.unload_model()
                self.current_model = selected_model
                self.update_model_path()
                self.update_model_label()
            else:
                # Revert to previous model selection
                if self.current_model == "4B":
                    self.model_combo.setCurrentIndex(0)
                else:
                    self.model_combo.setCurrentIndex(1)
        else:
            # Model chưa được load, chỉ cần cập nhật path
            self.current_model = selected_model
            self.update_model_path()
            self.update_model_label()
    
    def update_model_path(self):
        """Update model path based on current model selection"""
        if self.current_model == "2B":
            self.local_model_path = "./models/Qwen3-VL-2B-Instruct"
        else:
            self.local_model_path = "./models/Qwen3-VL-4B-Instruct"
    
    def update_model_label(self):
        """Update model label based on current model selection"""
        if self.current_model == "2B":
            model_name = "Model OCR nhẹ"
            model_info = "Nhẹ hơn (~4GB), nhanh hơn"
        else:
            model_name = "Model OCR mặc định"
            model_info = "Chính xác hơn (~8GB)"
        
        self.model_label.setText(f"Model: {model_name} ({model_info})")
        
        # Check if model exists
        if os.path.exists(self.local_model_path):
            self.model_label.setStyleSheet("color: #27ae60; font-size: 9pt; font-weight: bold; padding: 5px;")
        else:
            self.model_label.setStyleSheet("color: #e74c3c; font-size: 9pt; font-weight: bold; padding: 5px;")
            self.model_label.setText(f"Model: {model_name} ({model_info}) - Chưa tải về")
    
    def on_prompt_changed(self, text):
        # Khi chọn prompt mặc định, có thể điền vào custom prompt
        pass
    
    def on_use_default_prompt(self):
        """Điền prompt mặc định vào custom prompt input"""
        default_prompt = self.prompt_combo.currentText()
        self.custom_prompt.setPlainText(default_prompt)
    
    def on_auto_load_changed(self, state):
        """Handle auto-load checkbox change"""
        self.auto_load_enabled = (state == Qt.CheckState.Checked.value)
    
    def check_system_requirements(self):
        """Check if system meets requirements for auto-loading"""
        requirements = {
            "model_exists": False,
            "enough_memory": False,
            "gpu_available": False,
            "errors": []
        }
        
        # Check 1: Model files exist
        if os.path.exists(self.local_model_path):
            # Check key files - hỗ trợ cả model 2B (single file) và 4B (sharded)
            key_files = ["config.json"]
            
            # Model 2B có file model.safetensors duy nhất, model 4B có model.safetensors.index.json
            model_2b_file = os.path.join(self.local_model_path, "model.safetensors")
            model_4b_index = os.path.join(self.local_model_path, "model.safetensors.index.json")
            
            all_files_exist = True
            for file in key_files:
                if not os.path.exists(os.path.join(self.local_model_path, file)):
                    all_files_exist = False
                    break
            
            # Check ít nhất một trong hai: model.safetensors (2B) hoặc model.safetensors.index.json (4B)
            if all_files_exist:
                if not os.path.exists(model_2b_file) and not os.path.exists(model_4b_index):
                    all_files_exist = False
            
            if all_files_exist:
                requirements["model_exists"] = True
            else:
                requirements["errors"].append("File model không đầy đủ")
        else:
            requirements["errors"].append(f"Không tìm thấy model tại {self.local_model_path}")
        
        # Check 2: GPU availability (if GPU mode selected)
        selected_device = self.device_combo.currentData()
        if selected_device == "cuda":
            torch = get_torch()
            if is_cuda_available():
                gpu_memory = torch.cuda.get_device_properties(0).total_memory / (1024**3)
                if gpu_memory >= 10:  # Need at least 10GB VRAM
                    requirements["gpu_available"] = True
                    requirements["enough_memory"] = True
                else:
                    requirements["errors"].append(f"GPU VRAM quá nhỏ: {gpu_memory:.1f}GB (cần 10GB+)")
            else:
                requirements["errors"].append("CUDA không khả dụng")
        
        # Check 3: RAM check (for CPU mode or loading process)
        if PSUTIL_AVAILABLE:
            try:
                ram = psutil.virtual_memory()
                ram_gb = ram.total / (1024**3)
                ram_available_gb = ram.available / (1024**3)
                
                if ram_gb >= 16 and ram_available_gb >= 8:
                    if not requirements["enough_memory"]:
                        requirements["enough_memory"] = True
                else:
                    requirements["errors"].append(
                        f"RAM không đủ: {ram_available_gb:.1f}GB khả dụng, {ram_gb:.1f}GB tổng (cần 16GB+ tổng, 8GB+ khả dụng)"
                    )
            except:
                requirements["errors"].append("Không thể kiểm tra RAM")
        else:
            # If psutil not available, skip RAM check
            requirements["enough_memory"] = True  # Assume OK if can't check
        
        return requirements
    
    def auto_load_model(self):
        """Auto-load model if system requirements are met"""
        try:
            if not self.auto_load_enabled:
                return
            
            # Check system requirements
            self.model_status.setText("Trạng thái: Đang kiểm tra yêu cầu hệ thống...")
            self.model_status.setStyleSheet("color: orange;")
            QApplication.processEvents()
            
            try:
                requirements = self.check_system_requirements()
            except Exception as e:
                self.model_status.setText("Trạng thái: Lỗi kiểm tra hệ thống - Tắt tự động tải")
                self.model_status.setStyleSheet("color: red;")
                self.output_text.setText(f"Lỗi khi kiểm tra hệ thống: {str(e)}\n\nBạn có thể tải model thủ công.")
                return
            
            if not requirements["model_exists"]:
                self.model_status.setText("Trạng thái: Không tìm thấy model - Cần tải thủ công")
                self.model_status.setStyleSheet("color: red;")
                return
            
            if requirements["errors"]:
                # Show errors but still allow manual load
                error_msg = "Tự động tải bị tắt:\n" + "\n".join(requirements["errors"][:2])
                self.model_status.setText(f"Trạng thái: {error_msg}")
                self.model_status.setStyleSheet("color: orange;")
                self.output_text.setText(
                    "Tự động tải bị bỏ qua do:\n" + "\n".join(requirements["errors"]) +
                    "\n\nBạn vẫn có thể tải model thủ công."
                )
                return
            
            # All requirements met - auto-load
            self.model_status.setText("Trạng thái: Đang tự động tải model (đã kiểm tra hệ thống)...")
            self.model_status.setStyleSheet("color: orange;")
            QApplication.processEvents()
            
            # Load model (but don't raise exception to UI)
            try:
                # Call load_model but catch exceptions
                self.load_model()
            except KeyboardInterrupt:
                # User interrupted - just stop
                self.model_status.setText("Trạng thái: Tự động tải bị hủy")
                self.model_status.setStyleSheet("color: orange;")
                return
            except SystemExit:
                # Don't catch system exit
                raise
            except Exception as e:
                error_str = str(e)
                self.model_status.setText("Trạng thái: Tự động tải thất bại - Có thể tải thủ công")
                self.model_status.setStyleSheet("color: red;")
                
                # Provide detailed error message
                if "paging file" in error_str.lower() or "1455" in error_str or "memory" in error_str.lower():
                    error_msg = (
                        f"Lỗi Memory khi tự động tải:\n{error_str}\n\n"
                        f"Giải pháp:\n"
                        f"1. Tăng Windows paging file (chạy: increase_paging_file.bat)\n"
                        f"2. Đóng các ứng dụng khác\n"
                        f"3. Khởi động lại máy\n"
                        f"4. Tắt tự động tải và tải thủ công sau\n\n"
                        f"Bạn vẫn có thể tải model thủ công từ nút 'Sử dụng AI này'."
                    )
                else:
                    error_msg = f"Lỗi tự động tải: {error_str}\n\nBạn có thể tải model thủ công."
                
                self.output_text.setText(error_msg)
        except KeyboardInterrupt:
            # User interrupted - restore UI state
            self.model_status.setText("Trạng thái: Tự động tải bị hủy")
            self.model_status.setStyleSheet("color: orange;")
        except SystemExit:
            # Re-raise system exit
            raise
        except Exception as e:
            # Catch-all for any unexpected errors
            try:
                self.model_status.setText("Trạng thái: Lỗi không mong đợi - Ứng dụng vẫn hoạt động")
                self.model_status.setStyleSheet("color: red;")
                self.output_text.setText(f"Lỗi không mong đợi khi tự động tải: {str(e)}\n\nBạn có thể tải model thủ công.")
            except:
                # If even UI update fails, just print to console
                print(f"Critical error in auto_load_model: {e}")
    
    def on_device_changed(self, index):
        """Handle device selection change - auto switch model based on device"""
        selected_device = self.device_combo.currentData()
        
        # Auto-switch model based on device
        if selected_device == "cpu":
            # CPU selected - recommend light model (2B)
            if self.current_model == "4B":
                # Show notification and switch to 2B
                from PyQt6.QtWidgets import QMessageBox
                msg = QMessageBox(self)
                msg.setIcon(QMessageBox.Icon.Information)
                msg.setWindowTitle("Chuyển Model")
                msg.setText("CPU đã được chọn - Tự động chuyển sang Model OCR nhẹ (2B)")
                msg.setInformativeText("Model nhẹ (2B) chạy tốt hơn trên CPU.\n\nModel mạnh (4B) rất nặng và chậm trên CPU.")
                msg.setStandardButtons(QMessageBox.StandardButton.Ok)
                msg.exec()
                
                # Switch to 2B model
                self.model_combo.setCurrentIndex(1)  # Index 1 = 2B
                
        elif selected_device == "cuda":
            # GPU selected - can use strong model (4B)
            if self.current_model == "2B":
                # Ask if want to switch to 4B
                from PyQt6.QtWidgets import QMessageBox
                msg = QMessageBox(self)
                msg.setIcon(QMessageBox.Icon.Question)
                msg.setWindowTitle("Chuyển Model")
                msg.setText("GPU đã được chọn - Bạn có muốn chuyển sang Model OCR mạnh (4B)?")
                msg.setInformativeText("Model mạnh (4B) chính xác hơn và chạy tốt trên GPU.\n\nNhấn Yes để chuyển, No để giữ model nhẹ.")
                msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                msg.setDefaultButton(QMessageBox.StandardButton.Yes)
                
                if msg.exec() == QMessageBox.StandardButton.Yes:
                    # Switch to 4B model
                    self.model_combo.setCurrentIndex(0)  # Index 0 = 4B
        
        # Continue with original logic
        # Cập nhật label khuyến nghị dựa trên thiết bị đã chọn
        selected_device = self.device_combo.currentData()
        
        cuda_info = get_cuda_info()
        if cuda_info['available']:
            if selected_device == "cuda":
                # Người dùng chọn GPU - hiển thị khuyến nghị
                self.device_recommendation.setText("Khuyến nghị: Sử dụng GPU để có kết quả nhanh và tốt hơn CPU")
                self.device_recommendation.setStyleSheet("color: green; font-size: 9pt; font-weight: bold; padding: 5px;")
                self.device_combo.setToolTip("Khuyến nghị: Chọn GPU để có kết quả nhanh và tốt hơn CPU")
            else:
                # Người dùng chọn CPU - hiển thị cảnh báo
                self.device_recommendation.setText("Cảnh báo: CPU chậm hơn GPU đáng kể. Khuyến nghị sử dụng GPU để có kết quả tốt hơn.")
                self.device_recommendation.setStyleSheet("color: orange; font-size: 9pt; font-weight: bold; padding: 5px;")
                self.device_combo.setToolTip("Cảnh báo: CPU chậm hơn GPU đáng kể. Khuyến nghị chọn GPU để có kết quả tốt hơn.")
        else:
            # Không có GPU - chỉ có CPU
            self.device_recommendation.setText("Không có GPU - Sẽ sử dụng CPU (chậm hơn)")
            self.device_recommendation.setStyleSheet("color: orange; font-size: 9pt; font-weight: bold; padding: 5px;")
            self.device_combo.setToolTip("Không có GPU - Sẽ sử dụng CPU (chậm hơn GPU)")
        
        if self.model is not None:
            # Save current selection temporarily
            new_device = selected_device
            
            # Model is already loaded, need to reload with new device
            reply = self.show_device_change_warning()
            if reply:
                self.unload_model()
                self.load_model()
            else:
                # Revert to previous device selection
                torch = get_torch()
                if is_cuda_available():
                    if self.current_device == "cuda":
                        self.device_combo.setCurrentIndex(0)
                    else:
                        self.device_combo.setCurrentIndex(1)
                else:
                    self.device_combo.setCurrentIndex(0)
    
    def show_device_change_warning(self):
        """Show warning dialog when changing device"""
        from PyQt6.QtWidgets import QMessageBox
        msg = QMessageBox(self)
        msg.setIcon(QMessageBox.Icon.Warning)
        msg.setWindowTitle("Thay Đổi Thiết Bị")
        msg.setText("Model đã được tải. Thay đổi thiết bị cần tải lại model.")
        msg.setInformativeText("Bạn có muốn tải lại model với thiết bị mới không?")
        msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        msg.setDefaultButton(QMessageBox.StandardButton.Yes)
        return msg.exec() == QMessageBox.StandardButton.Yes
    
    def unload_model(self):
        """Unload model from memory"""
        if self.model is not None:
            del self.model
            self.model = None
        if self.processor is not None:
            del self.processor
            self.processor = None
        
        # Clear CUDA cache if available
        torch = get_torch()
        if is_cuda_available():
            torch.cuda.empty_cache()
        
        self.model_status.setText("Trạng thái: Model đã được gỡ")
        self.model_status.setStyleSheet("color: orange;")
        self.load_model_btn.setEnabled(True)
        self.unload_model_btn.setEnabled(False)
        self.update_process_button_state()
    
    def load_model(self):
        """Load model from local storage only using worker thread"""
        # Check if local model exists
        if not os.path.exists(self.local_model_path):
            self.model_status.setText("Trạng thái: Không tìm thấy model local!")
            self.model_status.setStyleSheet("color: red;")
            self.output_text.setText(
                f"Lỗi: Không tìm thấy model local tại {os.path.abspath(self.local_model_path)}\n\n"
                "Vui lòng chạy 'python download_model.py' để tải model trước."
            )
            self.load_model_btn.setEnabled(True)
            return
        
        # Don't allow multiple load operations
        if self.model_loader_worker is not None and self.model_loader_worker.isRunning():
            return
        
        # Create progress dialog
        self.progress_dialog = QProgressDialog("Đang tải model...", "Hủy", 0, 0, self)
        self.progress_dialog.setWindowTitle("Sử dụng AI này")
        self.progress_dialog.setWindowModality(Qt.WindowModality.WindowModal)
        self.progress_dialog.setAutoClose(False)
        self.progress_dialog.setAutoReset(False)
        self.progress_dialog.setMinimumDuration(0)  # Show immediately
        self.progress_dialog.setCancelButton(None)  # Disable cancel button
        self.progress_dialog.setRange(0, 0)  # Indeterminate progress
        self.progress_dialog.show()
        QApplication.processEvents()
        
        self.model_status.setText("Trạng thái: Đang tải model từ local storage...")
        self.model_status.setStyleSheet("color: orange;")
        self.load_model_btn.setEnabled(False)
        self.model_combo.setEnabled(False)  # Disable model selection khi đang load
        self.device_combo.setEnabled(False)
        QApplication.processEvents()
        
        # Get selected device
        selected_device = self.device_combo.currentData()
        self.current_device = selected_device
        
        # Determine device_map and dtype based on selection
        torch = get_torch()
        if selected_device == "cuda" and is_cuda_available():
            gpu_memory = torch.cuda.get_device_properties(0).total_memory / (1024**3)  # GB
            device_map = "cuda:0"
            dtype = torch.float16
            max_memory = {0: f"{int(gpu_memory)}GiB"}
            low_cpu_mem_usage = True
            
            # Show GPU info
            gpu_free = (torch.cuda.get_device_properties(0).total_memory - torch.cuda.memory_reserved(0)) / (1024**3)
            status_text = (
                f"Đang tải model vào GPU VRAM...\n"
                f"GPU: {torch.cuda.get_device_name(0)}\n"
                f"VRAM: {gpu_free:.1f}GB trống / {gpu_memory:.1f}GB tổng\n"
                f"Có thể mất 1-3 phút, vui lòng đợi..."
            )
        else:
            device_map = "cpu"
            dtype = torch.float32
            max_memory = None
            low_cpu_mem_usage = True
            status_text = (
                f"Đang tải model vào RAM (chế độ CPU)...\n"
                f"Có thể mất 2-5 phút, vui lòng đợi..."
            )
        
        self.model_status.setText(f"Trạng thái: {status_text}")
        self.progress_dialog.setLabelText(status_text)
        QApplication.processEvents()
        
        # Disable model selection và device selection khi đang load
        self.model_combo.setEnabled(False)
        self.device_combo.setEnabled(False)
        
        # Create and start worker thread
        self.model_loader_worker = ModelLoaderWorker(
            self.local_model_path,
            device_map,
            dtype,
            low_cpu_mem_usage,
            max_memory
        )
        self.model_loader_worker.finished.connect(self.on_model_load_finished)
        self.model_loader_worker.error.connect(self.on_model_load_error)
        self.model_loader_worker.progress.connect(self.on_model_load_progress)
        self.model_loader_worker.start()
    
    def on_model_load_progress(self, message):
        """Update progress dialog with loading status"""
        if self.progress_dialog is not None and message:
            self.progress_dialog.setLabelText(message)
            QApplication.processEvents()
    
    def on_model_load_finished(self, model, processor):
        """Handle successful model loading"""
        try:
            # Close progress dialog
            if self.progress_dialog is not None:
                self.progress_dialog.close()
                self.progress_dialog = None
            
            # Set model and processor
            self.model = model
            self.processor = processor
            
            # Get selected device for info display
            selected_device = self.device_combo.currentData()
            
            # Verify model is on correct device
            if hasattr(self.model, 'device'):
                actual_device = str(self.model.device)
            else:
                actual_device = str(next(self.model.parameters()).device)
            
            # Kiểm tra device thực tế có khớp với device đã chọn không
            device_match = False
            torch = get_torch()
            if selected_device == "cuda" and is_cuda_available():
                device_match = "cuda" in actual_device.lower()
                gpu_used = torch.cuda.memory_reserved(0) / (1024**3)
                gpu_total = torch.cuda.get_device_properties(0).total_memory / (1024**3)
                gpu_name = torch.cuda.get_device_name(0)
                device_info = f"GPU ({gpu_name})"
                memory_info = f"VRAM đã dùng: {gpu_used:.1f}GB / {gpu_total:.1f}GB"
                
                if not device_match:
                    print(f"[WARNING] Model được load vào {actual_device} nhưng đã chọn GPU!")
            else:
                device_match = "cpu" in actual_device.lower()
                device_info = "CPU"
                memory_info = "Chế độ RAM (CPU)"
                
                if not device_match:
                    print(f"[WARNING] Model được load vào {actual_device} nhưng đã chọn CPU!")
                    print(f"[WARNING] Model sẽ chạy chậm vì không ở CPU RAM!")
            
            # Hiển thị thông báo với icon tương ứng
            status_prefix = "" if device_match else "[CẢNH BÁO] "
            
            self.model_status.setText(
                f"{status_prefix}Trạng thái: Model đã tải từ local trên {device_info}\n"
                f"Device thực tế: {actual_device}\n"
                f"{memory_info}\n"
                f"Chế độ tối ưu memory: BẬT"
            )
            
            if not device_match:
                self.model_status.setStyleSheet("""
                    color: #e74c3c;
                    font-size: 10pt;
                    font-weight: bold;
                    padding: 8px;
                    background-color: #ffe6e6;
                    border: 2px solid #e74c3c;
                    border-radius: 8px;
                """)
            else:
                self.model_status.setStyleSheet("""
                    color: #27ae60;
                    font-size: 10pt;
                    font-weight: bold;
                    padding: 8px;
                    background-color: white;
                    border: 2px solid #e1e8ed;
                    border-radius: 8px;
                """)
            
            self.load_model_btn.setEnabled(True)
            self.unload_model_btn.setEnabled(True)
            self.update_process_button_state()
            self.device_combo.setEnabled(True)
            self.model_combo.setEnabled(True)  # Enable model selection sau khi load xong
            
            # Clean up worker
            self.model_loader_worker = None
            
        except Exception as e:
            self.on_model_load_error(str(e))
    
    def on_model_load_error(self, error_msg):
        """Handle model loading error - hiển thị popup"""
        import traceback
        
        # Close progress dialog if it exists
        if self.progress_dialog is not None:
            try:
                self.progress_dialog.close()
            except:
                pass
            self.progress_dialog = None
        
        print(f"\n[Main Thread] Model load error received: {error_msg}")
        
        self.model_status.setText(f"Trạng thái: Lỗi khi tải model")
        self.model_status.setStyleSheet("color: red;")
        
        # Provide helpful error message for memory issues
        if "paging file" in error_msg.lower() or "paging file quá nhỏ" in error_msg.lower():
            # Nếu là lỗi paging file, hiển thị thông báo đặc biệt
            detailed_msg = error_msg  # Dùng error_msg từ worker thread (đã có hướng dẫn chi tiết)
            detailed_msg += f"\n\nTIP: Chạy 'increase_paging_file.bat' để mở System Properties nhanh hơn!"
        elif "1455" in error_msg or "memory" in error_msg.lower() or "out of memory" in error_msg.lower():
            detailed_msg = (
                f"Lỗi Memory: Không đủ RAM/virtual memory để tải model.\n\n"
                f"Lỗi: {error_msg}\n\n"
                f"Giải pháp:\n"
                f"1. Đóng các ứng dụng khác để giải phóng RAM\n"
                f"2. Tăng Windows virtual memory (paging file) lên ít nhất 8GB (chạy: increase_paging_file.bat)\n"
                f"3. RESTART MÁY sau khi tăng paging file\n"
                f"4. Thử tải trên CPU thay vì GPU (ít dùng memory hơn)\n"
                f"5. Hệ thống cần ít nhất 16GB RAM (khuyến nghị 32GB)\n\n"
                f"Đường dẫn model: {os.path.abspath(self.local_model_path)}\n"
            )
        else:
            detailed_msg = (
                f"Lỗi khi tải model: {error_msg}\n\n"
                f"Đường dẫn model: {os.path.abspath(self.local_model_path)}\n"
                f"Hãy đảm bảo model đã được tải về đúng cách."
            )
        
        self.output_text.setText(detailed_msg)
        self.load_model_btn.setEnabled(True)
        self.device_combo.setEnabled(True)
        self.model_combo.setEnabled(True)  # Enable model selection sau khi load lỗi
        
        # Show error popup
        try:
            msg = QMessageBox(self)
            msg.setIcon(QMessageBox.Icon.Critical)
            msg.setWindowTitle("Lỗi Khi Sử Dụng AI")
            msg.setText("Không thể tải model!\n\nVui lòng xem chi tiết bên dưới.")
            msg.setDetailedText(error_msg)
            msg.setStandardButtons(QMessageBox.StandardButton.Ok)
            msg.setDefaultButton(QMessageBox.StandardButton.Ok)
            msg.exec()
        except Exception as e:
            print(f"[Main Thread] Không thể hiển thị popup: {e}")
        
        # Clean up worker
        self.model_loader_worker = None
    
    def load_file(self):
        """Load file (Image, PDF, DOCX, TXT)"""
        file_name, _ = QFileDialog.getOpenFileName(
            self,
            "Mở File",
            "",
            "Tất Cả Định Dạng Hỗ Trợ (*.png *.jpg *.jpeg *.bmp *.gif *.pdf *.docx *.txt);;"
            "File Hình Ảnh (*.png *.jpg *.jpeg *.bmp *.gif);;"
            "File PDF (*.pdf);;"
            "File DOCX (*.docx);;"
            "File Text (*.txt);;"
            "Tất Cả File (*)"
        )
        
        if not file_name:
            return
        
        # Clean up previous temp files
        self.cleanup_temp_files()
        
        file_ext = os.path.splitext(file_name)[1].lower()
        
        try:
            if file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
                self.load_image_file(file_name)
            elif file_ext == '.pdf':
                self.load_pdf_file(file_name)
            elif file_ext == '.docx':
                self.load_docx_file(file_name)
            elif file_ext == '.txt':
                self.load_txt_file(file_name)
            else:
                QMessageBox.warning(
                    self, 
                    "File Không Được Hỗ Trợ",
                    f"Loại file '{file_ext}' không được hỗ trợ.\n\n"
                    "Định dạng được hỗ trợ: JPG, PNG, JPEG, PDF, DOCX, TXT"
                )
                return
            
            self.current_file_path = file_name
            self.update_process_button_state()
            
        except Exception as e:
            QMessageBox.critical(
                self,
                "Lỗi Khi Tải File",
                f"Lỗi khi tải file: {str(e)}\n\nVui lòng kiểm tra file có hợp lệ không."
            )
    
    def load_multiple_files(self):
        """Load multiple files for batch processing"""
        if not self.model or not self.processor:
            QMessageBox.warning(
                self,
                "Chưa Tải Model",
                "Vui lòng tải model trước khi xử lý nhiều file!\n\nClick nút 'Sử dụng AI này' để tải model."
            )
            return
        
        file_names, _ = QFileDialog.getOpenFileNames(
            self,
            "Chọn Nhiều File",
            "",
            "Tất Cả Định Dạng Hỗ Trợ (*.png *.jpg *.jpeg *.bmp *.gif *.pdf *.docx *.txt);;"
            "File Hình Ảnh (*.png *.jpg *.jpeg *.bmp *.gif);;"
            "File PDF (*.pdf);;"
            "File DOCX (*.docx);;"
            "File Text (*.txt);;"
            "Tất Cả File (*)"
        )
        
        if not file_names:
            return
        
        # Prepare file list for batch processing
        file_list = []
        temp_files_to_clean = []
        
        try:
            for file_path in file_names:
                file_ext = os.path.splitext(file_path)[1].lower()
                
                if file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
                    file_list.append((file_path, 'image', file_path))
                
                elif file_ext == '.pdf':
                    if not PDF_SUPPORT:
                        continue
                    try:
                        from pdf2image import convert_from_path
                        images = convert_from_path(file_path, first_page=1, last_page=1, dpi=200)
                        if images:
                            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                            images[0].save(temp_file.name, 'PNG')
                            temp_files_to_clean.append(temp_file.name)
                            file_list.append((file_path, 'pdf', temp_file.name))
                    except Exception as e:
                        print(f"[Batch] Lỗi khi xử lý PDF {file_path}: {e}")
                        continue
                
                elif file_ext == '.docx':
                    if not DOCX_SUPPORT:
                        continue
                    try:
                        from docx import Document
                        from PIL import Image, ImageDraw, ImageFont
                        doc = Document(file_path)
                        text_content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                        if text_content:
                            img = Image.new('RGB', (800, min(600, len(text_content.split('\n')) * 25 + 40)), color='white')
                            draw = ImageDraw.Draw(img)
                            try:
                                font = ImageFont.truetype("arial.ttf", 16)
                            except:
                                font = ImageFont.load_default()
                            y = 20
                            for line in text_content.split('\n')[:50]:
                                draw.text((20, y), line[:120], fill='black', font=font)
                                y += 25
                            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                            img.save(temp_file.name, 'PNG')
                            temp_files_to_clean.append(temp_file.name)
                            file_list.append((file_path, 'docx', temp_file.name))
                    except Exception as e:
                        print(f"[Batch] Lỗi khi xử lý DOCX {file_path}: {e}")
                        continue
                
                elif file_ext == '.txt':
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                        if content.strip():
                            from PIL import Image, ImageDraw, ImageFont
                            lines = content.split('\n')
                            img = Image.new('RGB', (800, min(600, len(lines) * 25 + 40)), color='white')
                            draw = ImageDraw.Draw(img)
                            try:
                                font = ImageFont.truetype("arial.ttf", 16)
                            except:
                                font = ImageFont.load_default()
                            y = 20
                            for line in lines[:50]:
                                draw.text((20, y), line[:120], fill='black', font=font)
                                y += 25
                            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                            img.save(temp_file.name, 'PNG')
                            temp_files_to_clean.append(temp_file.name)
                            file_list.append((file_path, 'txt', temp_file.name))
                    except Exception as e:
                        print(f"[Batch] Lỗi khi xử lý TXT {file_path}: {e}")
                        continue
            
            if not file_list:
                QMessageBox.warning(self, "Không Có File Hợp Lệ", "Không có file nào hợp lệ để xử lý!")
                return
            
            # Store temp files for cleanup
            self.batch_temp_files = temp_files_to_clean
            
            # Start batch processing
            self.start_batch_processing(file_list)
            
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi chuẩn bị batch processing: {str(e)}")
    
    def start_batch_processing(self, file_list):
        """Bắt đầu xử lý batch"""
        if not self.model or not self.processor:
            return
        
        self.batch_queue = file_list
        self.total_batch_files = len(file_list)
        self.current_batch_index = 0
        self.batch_results = []
        
        # Update UI
        self.batch_info_label.setText(f"Đã chọn {self.total_batch_files} file. Đang xử lý...")
        self.batch_progress_bar.setVisible(True)
        self.batch_progress_bar.setMaximum(self.total_batch_files)
        self.batch_progress_bar.setValue(0)
        self.load_multiple_files_btn.setEnabled(False)
        self.process_btn.setEnabled(False)
        
        # Get prompt and generation params
        prompt = self.get_prompt()
        generation_params = self.get_generation_params()
        
        # Create and start batch worker
        self.batch_worker = BatchWorker(
            self.model,
            self.processor,
            file_list,
            prompt,
            generation_params
        )
        self.batch_worker.progress.connect(self.on_batch_progress)
        self.batch_worker.file_finished.connect(self.on_batch_file_finished)
        self.batch_worker.file_error.connect(self.on_batch_file_error)
        self.batch_worker.finished.connect(self.on_batch_finished)
        self.batch_worker.start()
    
    def on_batch_progress(self, current, total, file_name):
        """Update batch progress"""
        self.current_batch_index = current
        self.batch_progress_bar.setValue(current)
        self.batch_info_label.setText(f"Đang xử lý: {current}/{total} - {file_name}")
    
    def on_batch_file_finished(self, file_path, file_type, ocr_result):
        """Xử lý khi một file xử lý xong"""
        # Lưu vào history
        processing_time = None  # Không track thời gian cho từng file trong batch
        self.save_history(file_path, file_type, ocr_result, processing_time)
        
        self.batch_results.append({
            'file_path': file_path,
            'file_type': file_type,
            'result': ocr_result,
            'success': True
        })
    
    def on_batch_file_error(self, file_path, error_message):
        """Xử lý khi một file bị lỗi"""
        print(f"[Batch] Lỗi: {error_message}")
        self.batch_results.append({
            'file_path': file_path,
            'file_type': 'unknown',
            'result': '',
            'success': False,
            'error': error_message
        })
    
    def on_batch_finished(self):
        """Xử lý khi batch processing hoàn thành"""
        # Cleanup temp files
        if hasattr(self, 'batch_temp_files'):
            for temp_file in self.batch_temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                except:
                    pass
            self.batch_temp_files = []
        
        # Update UI
        success_count = sum(1 for r in self.batch_results if r.get('success', False))
        error_count = len(self.batch_results) - success_count
        
        # Stop batch progress bar - ensure it's at 100% and stopped
        self.batch_progress_bar.setValue(self.total_batch_files)
        self.batch_progress_bar.setMaximum(self.total_batch_files)  # Ensure range is correct
        # Optional: hide after completion
        # self.batch_progress_bar.setVisible(False)
        
        self.batch_info_label.setText(
            f"Hoàn thành! Thành công: {success_count}/{self.total_batch_files}, "
            f"Lỗi: {error_count}"
        )
        self.load_multiple_files_btn.setEnabled(True)
        self.process_btn.setEnabled(True)
        
        # Show summary
        QMessageBox.information(
            self,
            "Hoàn Thành Batch Processing",
            f"Đã xử lý xong {self.total_batch_files} file:\n\n"
            f"Thành công: {success_count}\n"
            f"Lỗi: {error_count}\n\n"
            f"Tất cả kết quả đã được lưu vào History."
        )
        
        # Clear batch worker
        self.batch_worker = None
    
    def load_image_file(self, file_path):
        """Load image file"""
        self.current_file_type = 'image'
        self.current_image_path = file_path
        
        # Display image
        pixmap = QPixmap(file_path)
        if pixmap.isNull():
            raise ValueError("File hình ảnh không hợp lệ")
        
        # Set pixmap to custom ImageROILabel
        self.image_label.set_pixmap(pixmap)
        self.file_info_label.setText(f"File: {os.path.basename(file_path)} (Hình ảnh)")
        self.file_info_label.setStyleSheet("color: green; font-size: 9pt;")
        
        # Clear ROI when loading new image
        self.image_label.clear_roi()
    
    def load_pdf_file(self, file_path):
        """Load PDF file - convert first page to image"""
        if not PDF_SUPPORT:
            QMessageBox.warning(
                self,
                "Không Hỗ Trợ PDF",
                "Hỗ trợ PDF cần thư viện pdf2image.\n\n"
                "Cài đặt: pip install pdf2image\n"
                "Cũng cần poppler: https://github.com/oschwartz10612/poppler-windows/releases\n\n"
                "Hoặc chạy: python install_poppler.py"
            )
            return
        
        self.current_file_type = 'pdf'
        
        try:
            # Try with local poppler first (if installed via script)
            # Handle different ZIP structures: poppler-XX.XX.X/Library/bin or poppler/Library/bin
            poppler_path = None
            poppler_base = "./poppler"
            if os.path.exists(poppler_base):
                # Check for nested structure: poppler-XX.XX.X/Library/bin
                for item in os.listdir(poppler_base):
                    item_path = os.path.join(poppler_base, item)
                    if os.path.isdir(item_path) and item.startswith("poppler-"):
                        bin_path = os.path.join(item_path, "Library", "bin")
                        if os.path.exists(bin_path):
                            poppler_path = os.path.abspath(bin_path)
                            break
                
                # Check for direct structure: poppler/Library/bin
                if not poppler_path:
                    bin_path = os.path.join(poppler_base, "Library", "bin")
                    if os.path.exists(bin_path):
                        poppler_path = os.path.abspath(bin_path)
            
            # Convert first page to image
            if poppler_path and os.path.exists(os.path.join(poppler_path, "pdftoppm.exe")):
                # Use local poppler
                images = convert_from_path(
                    file_path, 
                    first_page=1, 
                    last_page=1, 
                    dpi=200,
                    poppler_path=poppler_path
                )
            else:
                # Try system PATH poppler
                images = convert_from_path(file_path, first_page=1, last_page=1, dpi=200)
            
            if not images:
                raise ValueError("Không thể trích xuất hình ảnh từ PDF. Hãy đảm bảo Poppler đã được cài đặt.")
            
            # Save to temp file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            images[0].save(temp_file.name, 'PNG')
            self.temp_image_files.append(temp_file.name)
            
            self.current_image_path = temp_file.name
            
            # Display image
            pixmap = QPixmap(temp_file.name)
            self.image_label.set_pixmap(pixmap)
            self.file_info_label.setText(f"File: {os.path.basename(file_path)} (PDF - Trang 1)")
            self.file_info_label.setStyleSheet("color: blue; font-size: 9pt;")
            
            # Clear ROI when loading new file
            self.image_label.clear_roi()
            
        except Exception as e:
            error_msg = str(e)
            if "poppler" in error_msg.lower() or "pdftoppm" in error_msg.lower():
                raise Exception(
                    f"Không tìm thấy Poppler. Vui lòng cài đặt Poppler:\n\n"
                    f"1. Chạy: python install_poppler.py\n"
                    f"2. Hoặc tải từ: https://github.com/oschwartz10612/poppler-windows/releases\n"
                    f"3. Thêm Poppler bin folder vào PATH\n\n"
                    f"Lỗi: {error_msg}"
                )
            else:
                raise Exception(f"Lỗi khi xử lý PDF: {error_msg}")
    
    def load_docx_file(self, file_path):
        """Load DOCX file - convert to image"""
        if not DOCX_SUPPORT:
            QMessageBox.warning(
                self,
                "Không Hỗ Trợ DOCX",
                "Hỗ trợ DOCX cần thư viện python-docx.\n\n"
                "Cài đặt: pip install python-docx"
            )
            return
        
        self.current_file_type = 'docx'
        
        try:
            # Read DOCX
            doc = Document(file_path)
            
            # Extract text (for now, we'll convert to image)
            # Create a simple text image
            text_content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            if not text_content.strip():
                raise ValueError("File DOCX có vẻ trống")
            
            # Create image from text
            from PIL import ImageDraw, ImageFont
            
            # Estimate image size
            lines = text_content.split('\n')
            max_width = max(len(line) for line in lines) if lines else 80
            img_width = min(max_width * 10, 1200)
            img_height = len(lines) * 25 + 40
            
            # Create image
            img = Image.new('RGB', (img_width, img_height), color='white')
            draw = ImageDraw.Draw(img)
            
            # Try to use default font
            try:
                font = ImageFont.truetype("arial.ttf", 16)
            except:
                font = ImageFont.load_default()
            
            # Draw text
            y = 20
            for line in lines[:50]:  # Limit to 50 lines
                draw.text((20, y), line[:120], fill='black', font=font)
                y += 25
            
            # Save to temp file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            img.save(temp_file.name, 'PNG')
            self.temp_image_files.append(temp_file.name)
            
            self.current_image_path = temp_file.name
            
            # Display image
            pixmap = QPixmap(temp_file.name)
            self.image_label.set_pixmap(pixmap)
            self.file_info_label.setText(f"File: {os.path.basename(file_path)} (DOCX)")
            self.file_info_label.setStyleSheet("color: purple; font-size: 9pt;")
            
            # Clear ROI when loading new file
            self.image_label.clear_roi()
            
        except Exception as e:
            raise Exception(f"Lỗi khi xử lý DOCX: {str(e)}")
    
    def load_txt_file(self, file_path):
        """Load TXT file - read directly and show preview"""
        self.current_file_type = 'txt'
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if not content.strip():
                raise ValueError("File TXT có vẻ trống")
            
            # Create image from text for OCR
            from PIL import ImageDraw, ImageFont
            
            lines = content.split('\n')
            max_width = max(len(line) for line in lines) if lines else 80
            img_width = min(max_width * 10, 1200)
            img_height = min(len(lines) * 25 + 40, 3000)
            
            # Create image
            img = Image.new('RGB', (img_width, img_height), color='white')
            draw = ImageDraw.Draw(img)
            
            # Try to use default font
            try:
                font = ImageFont.truetype("arial.ttf", 16)
            except:
                font = ImageFont.load_default()
            
            # Draw text
            y = 20
            for line in lines[:100]:  # Limit to 100 lines
                draw.text((20, y), line[:120], fill='black', font=font)
                y += 25
            
            # Save to temp file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            img.save(temp_file.name, 'PNG')
            self.temp_image_files.append(temp_file.name)
            
            self.current_image_path = temp_file.name
            
            # Display image
            pixmap = QPixmap(temp_file.name)
            self.image_label.set_pixmap(pixmap)
            self.file_info_label.setText(f"File: {os.path.basename(file_path)} (TXT)")
            self.file_info_label.setStyleSheet("color: orange; font-size: 9pt;")
            
            # Clear ROI when loading new file
            self.image_label.clear_roi()
            
        except Exception as e:
            raise Exception(f"Lỗi khi xử lý TXT: {str(e)}")
    
    def cleanup_temp_files(self):
        """Clean up temporary image files"""
        for temp_file in self.temp_image_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except:
                pass
        self.temp_image_files = []
    
    def update_process_button_state(self):
        if self.model is not None and self.current_image_path is not None:
            self.process_btn.setEnabled(True)
        else:
            self.process_btn.setEnabled(False)
    
    def get_generation_params(self):
        return {
            "max_new_tokens": self.max_tokens_spin.value(),
            "temperature": self.temperature_spin.value(),
            "top_p": self.top_p_spin.value(),
            "top_k": self.top_k_spin.value(),
            "repetition_penalty": self.rep_penalty_spin.value(),
        }
    
    def get_prompt(self):
        # Ưu tiên custom prompt nếu có nội dung, nếu không thì dùng prompt mặc định
        custom_prompt_text = self.custom_prompt.toPlainText().strip()
        if custom_prompt_text:
            return custom_prompt_text
        return self.prompt_combo.currentText()
    
    def crop_image_by_roi(self, image_path, roi_rect):
        """Crop image theo ROI rectangle"""
        try:
            from PIL import Image
            
            # Load image
            img = Image.open(image_path).convert("RGB")
            
            # Crop theo ROI
            cropped = img.crop((
                roi_rect.x(),
                roi_rect.y(),
                roi_rect.x() + roi_rect.width(),
                roi_rect.y() + roi_rect.height()
            ))
            
            # Save to temp file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            cropped.save(temp_file.name, 'PNG')
            
            return temp_file.name
        except Exception as e:
            print(f"[ROI] Lỗi khi crop image: {e}")
            return None
    
    def get_image_for_processing(self):
        """Lấy image path để xử lý (có thể là cropped ROI hoặc full image)"""
        # Kiểm tra ROI mode và có ROI được chọn không
        if self.roi_mode_checkbox.isChecked():
            roi_rect = self.image_label.get_roi_rect()
            if roi_rect and not roi_rect.isEmpty():
                # Crop image theo ROI
                cropped_path = self.crop_image_by_roi(self.current_image_path, roi_rect)
                if cropped_path:
                    return cropped_path
        
        # Mặc định: dùng full image
        return self.current_image_path
    
    def process_image(self):
        if not self.model or not self.current_image_path:
            return
        
        # Sử dụng prompt duy nhất cho tất cả file types
        prompt = self.get_prompt()
        if not prompt or prompt.strip() == "":
            self.output_text.setText("Lỗi: Vui lòng nhập prompt")
            return
        
        # Ghi lại thời gian bắt đầu xử lý
        self.processing_start_time = time.time()
        self.processing_time_label.setText("Thời gian xử lý: Đang xử lý...")
        self.processing_time_label.setStyleSheet("color: orange; font-size: 10pt; font-weight: bold;")
        
        # Disable button during processing
        self.process_btn.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 0)  # Indeterminate progress
        self.output_text.clear()
        
        # Kiểm tra ROI mode
        if self.roi_mode_checkbox.isChecked():
            roi_rect = self.image_label.get_roi_rect()
            if not roi_rect or roi_rect.isEmpty():
                QMessageBox.warning(
                    self,
                    "Chưa Chọn Vùng",
                    "Vui lòng chọn vùng cần OCR bằng cách drag chuột trên hình ảnh!"
                )
                self.process_btn.setEnabled(True)
                return
        
        # Get image for processing (ROI cropped or full)
        image_to_process = self.get_image_for_processing()
        if not image_to_process:
            QMessageBox.warning(self, "Lỗi", "Không thể chuẩn bị hình ảnh để xử lý!")
            self.process_btn.setEnabled(True)
            return
        
        # Track temp file nếu có (cropped ROI)
        self.current_processing_image = image_to_process
        if image_to_process != self.current_image_path:
            # Cropped image, cần cleanup sau
            if not hasattr(self, 'processing_temp_files'):
                self.processing_temp_files = []
            self.processing_temp_files.append(image_to_process)
        
        # Get generation parameters
        gen_params = self.get_generation_params()
        
        # Create and start worker thread
        self.worker = ModelWorker(
            self.model,
            self.processor,
            image_to_process,
            prompt,
            gen_params
        )
        self.worker.finished.connect(self.on_processing_finished)
        self.worker.error.connect(self.on_processing_error)
        self.worker.progress.connect(self.on_progress_update)
        self.worker.start()
    
    def on_progress_update(self, message):
        self.progress_label.setText(message)
    
    def on_processing_finished(self, result):
        """Xử lý khi OCR hoàn thành"""
        # Calculate processing time
        elapsed_time = None
        if self.processing_start_time:
            elapsed_time = time.time() - self.processing_start_time
            hours = int(elapsed_time // 3600)
            minutes = int((elapsed_time % 3600) // 60)
            seconds = elapsed_time % 60
            
            if hours > 0:
                time_str = f"{hours}h {minutes}m {seconds:.1f}s"
            elif minutes > 0:
                time_str = f"{minutes}m {seconds:.1f}s"
            else:
                time_str = f"{seconds:.2f}s"
            
            self.processing_time_label.setText(f"Thời gian xử lý: {time_str}")
            self.processing_time_label.setStyleSheet("color: #27ae60; font-size: 11pt; font-weight: bold; padding: 8px; background-color: white; border-radius: 8px;")
        else:
            self.processing_time_label.setText("Thời gian xử lý: Hoàn thành")
            self.processing_time_label.setStyleSheet("color: #27ae60; font-size: 11pt; font-weight: bold; padding: 8px; background-color: white; border-radius: 8px;")
        
        # Display result - format if JSON
        formatted_result = self.format_result_display(result)
        self.output_text.setText(formatted_result)
        # Stop progress bar - set range and value to stop animation
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(100)
        # Hide progress bar after a short delay (optional, or keep it visible at 100%)
        # self.progress_bar.setVisible(False)  # Uncomment to hide completely
        self.progress_label.setText("Hoàn thành!")
        self.progress_label.setStyleSheet("color: #27ae60; font-size: 10pt;")
        
        # Re-enable button
        self.process_btn.setEnabled(True)
        
        # Auto-save to database
        if self.current_file_path and result:
            try:
                self.save_history(
                    file_path=self.current_file_path,
                    file_type=self.current_file_type or 'image',
                    ocr_result=result,
                    processing_time=elapsed_time
                )
                print(f"[Auto-Save] Đã tự động lưu kết quả OCR vào database")
            except Exception as save_error:
                print(f"[Auto-Save] Lỗi khi tự động lưu: {save_error}")
        
        self.processing_start_time = None
        
        # Cleanup temp files từ ROI crop
        if hasattr(self, 'processing_temp_files'):
            for temp_file in self.processing_temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                except:
                    pass
            self.processing_temp_files = []
    
    def on_roi_mode_changed(self, state):
        """Xử lý khi ROI mode checkbox thay đổi"""
        enabled = (state == Qt.CheckState.Checked.value)
        self.image_label.set_roi_mode(enabled)
        self.clear_roi_btn.setEnabled(enabled)
        
        if not enabled:
            self.roi_info_label.setVisible(False)
            self.image_label.clear_roi()
        else:
            self.roi_info_label.setVisible(True)
            self.roi_info_label.setText("Click và drag trên hình ảnh để chọn vùng OCR")
    
    def on_roi_changed(self, roi_rect):
        """Xử lý khi ROI rectangle thay đổi"""
        if roi_rect.isEmpty():
            self.roi_info_label.setText("Chưa chọn vùng. Click và drag để chọn vùng OCR")
        else:
            width = roi_rect.width()
            height = roi_rect.height()
            self.roi_info_label.setText(
                f"Đã chọn vùng: {width}x{height} pixels tại ({roi_rect.x()}, {roi_rect.y()})"
            )
    
    def clear_roi_selection(self):
        """Xóa vùng ROI đã chọn"""
        self.image_label.clear_roi()
        self.roi_info_label.setText("Chưa chọn vùng. Click và drag để chọn vùng OCR")
    
    def format_result_display(self, result):
        """Format kết quả để hiển thị đẹp hơn nếu là JSON"""
        try:
            # Try to parse JSON from result
            if '{' in result and '}' in result:
                json_start = result.find('{')
                json_end = result.rfind('}') + 1
                json_str = result[json_start:json_end]
                json_response = json.loads(json_str)
                
                # Format theo template với emoji
                formatted = "📄 KẾT QUẢ PHÂN TÍCH VĂN BẢN\n"
                formatted += "=" * 60 + "\n\n"
                formatted += f"🔹 LOẠI VĂN BẢN: {json_response.get('loai_van_ban', 'N/A')}\n\n"
                formatted += f"🔹 SỐ KÝ HIỆU: {json_response.get('so_ky_hieu', 'N/A')}\n\n"
                formatted += f"🔹 CƠ QUAN BAN HÀNH: {json_response.get('co_quan_ban_hanh', 'N/A')}\n\n"
                formatted += f"🔹 NGÀY BAN HÀNH: {json_response.get('ngay_ban_hanh', 'N/A')}\n\n"
                formatted += f"🔹 TRÍCH YẾU:\n{json_response.get('trich_yeu', 'N/A')}\n\n"
                formatted += "=" * 60 + "\n\n"
                formatted += "📋 JSON GỐC:\n"
                formatted += json.dumps(json_response, ensure_ascii=False, indent=2)
                
                return formatted
        except Exception as e:
            print(f"[Format] Không thể format kết quả: {e}")
        
        # Return original result if formatting fails
        return result
    
    def on_processing_error(self, error_message):
        # Cleanup temp files từ ROI crop
        if hasattr(self, 'processing_temp_files'):
            for temp_file in self.processing_temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                except:
                    pass
            self.processing_temp_files = []
        
        # Tính thời gian xử lý (nếu có)
        if self.processing_start_time is not None:
            elapsed_time = time.time() - self.processing_start_time
            
            # Định dạng thời gian: giây, phút, giờ
            if elapsed_time < 60:
                time_str = f"{elapsed_time:.2f} giây"
            elif elapsed_time < 3600:
                minutes = int(elapsed_time // 60)
                seconds = elapsed_time % 60
                time_str = f"{minutes} phút {seconds:.2f} giây"
            else:
                hours = int(elapsed_time // 3600)
                minutes = int((elapsed_time % 3600) // 60)
                seconds = elapsed_time % 60
                time_str = f"{hours} giờ {minutes} phút {seconds:.2f} giây"
            
            self.processing_time_label.setText(f"Thời gian xử lý (lỗi): {time_str}")
            self.processing_time_label.setStyleSheet("color: red; font-size: 10pt; font-weight: bold;")
            self.processing_start_time = None
        else:
            self.processing_time_label.setText("Thời gian xử lý: Lỗi xảy ra")
            self.processing_time_label.setStyleSheet("color: red; font-size: 10pt; font-weight: bold;")
        
        self.output_text.setText(error_message)
        # Stop progress bar - set range and value to stop animation
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)  # Set to 0 on error
        self.progress_bar.setVisible(False)
        self.progress_label.setText("Đã xảy ra lỗi")
        self.process_btn.setEnabled(True)
    
    def save_output(self):
        if not self.output_text.toPlainText():
            return
        
        file_name, _ = QFileDialog.getSaveFileName(
            self,
            "Lưu Kết Quả",
            "",
            "File JSON (*.json);;File Text (*.txt);;Tất Cả File (*)"
        )
        
        if file_name:
            try:
                result_text = self.output_text.toPlainText()
                
                # Kiểm tra extension để quyết định format
                if file_name.endswith('.json'):
                    # Try to parse existing JSON from result
                    json_response = None
                    try:
                        # Extract JSON from result text
                        if '{' in result_text and '}' in result_text:
                            json_start = result_text.find('{')
                            json_end = result_text.rfind('}') + 1
                            json_str = result_text[json_start:json_end]
                            json_response = json.loads(json_str)
                    except Exception as parse_error:
                        print(f"[Save] Không thể parse JSON từ kết quả: {parse_error}")
                    
                    # Prepare output data với format yêu cầu
                    if json_response and isinstance(json_response, dict):
                        # Format theo template với emoji và cấu trúc rõ ràng
                        output_data = {
                            "loai_van_ban": json_response.get('loai_van_ban', 'N/A'),
                            "so_ky_hieu": json_response.get('so_ky_hieu', 'N/A'),
                            "co_quan_ban_hanh": json_response.get('co_quan_ban_hanh', 'N/A'),
                            "ngay_ban_hanh": json_response.get('ngay_ban_hanh', 'N/A'),
                            "trich_yeu": json_response.get('trich_yeu', 'N/A'),
                            "metadata": {
                                "file_name": os.path.basename(self.current_file_path) if self.current_file_path else None,
                                "model": f"{self.current_model}B" if self.current_model else None,
                                "device": self.current_device,
                                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                "processing_time": self.processing_time_label.text()
                            }
                        }
                    else:
                        # Fallback: lưu raw result nếu không parse được JSON
                        output_data = {
                            "ocr_result": result_text,
                            "file_info": {
                                "file_path": self.current_file_path if self.current_file_path else None,
                                "file_name": os.path.basename(self.current_file_path) if self.current_file_path else None,
                                "file_type": self.current_file_type if self.current_file_type else None
                            },
                            "processing_info": {
                                "model": self.current_model,
                                "device": self.current_device,
                                "timestamp": datetime.now().isoformat(),
                                "processing_time": self.processing_time_label.text()
                            },
                            "generation_params": self.get_generation_params()
                        }
                    
                    # Lưu dạng JSON
                    with open(file_name, 'w', encoding='utf-8') as f:
                        json.dump(output_data, f, ensure_ascii=False, indent=2)
                else:
                    # Lưu dạng text thuần
                    with open(file_name, 'w', encoding='utf-8') as f:
                        f.write(result_text)
                
                self.progress_label.setText(f"Đã lưu kết quả vào {file_name}")
                self.progress_label.setStyleSheet("color: #27ae60; font-size: 10pt;")
            except Exception as e:
                self.progress_label.setText(f"Lỗi khi lưu file: {str(e)}")
                self.progress_label.setStyleSheet("color: red; font-size: 10pt;")


    def exit_fullscreen(self):
        """Exit fullscreen mode"""
        self.showNormal()
        # Update button visibility or text
        if hasattr(self, 'exit_fullscreen_btn'):
            self.exit_fullscreen_btn.setText("🗖 Fullscreen")
            self.exit_fullscreen_btn.disconnect()
            self.exit_fullscreen_btn.clicked.connect(self.enter_fullscreen)
    
    def enter_fullscreen(self):
        """Enter fullscreen mode"""
        self.showFullScreen()
        # Update button visibility or text
        if hasattr(self, 'exit_fullscreen_btn'):
            self.exit_fullscreen_btn.setText("🗗 Thoát Fullscreen")
            self.exit_fullscreen_btn.disconnect()
            self.exit_fullscreen_btn.clicked.connect(self.exit_fullscreen)
    
    def keyPressEvent(self, event):
        """Handle keyboard shortcuts"""
        from PyQt6.QtCore import Qt
        # ESC to exit fullscreen
        if event.key() == Qt.Key.Key_Escape:
            if self.isFullScreen():
                self.exit_fullscreen()
        # F11 to toggle fullscreen
        elif event.key() == Qt.Key.Key_F11:
            if self.isFullScreen():
                self.exit_fullscreen()
            else:
                self.enter_fullscreen()
        else:
            super().keyPressEvent(event)
    
    def closeEvent(self, event):
        """Clean up temp files on close"""
        self.cleanup_temp_files()
        event.accept()


def exception_hook(exc_type, exc_value, exc_traceback):
    """Global exception handler để hiển thị popup khi crash"""
    import traceback
    from PyQt6.QtWidgets import QMessageBox, QApplication
    from PyQt6.QtCore import Qt
    
    # Format error message
    error_msg = str(exc_value)
    traceback_str = ''.join(traceback.format_exception(exc_type, exc_value, exc_traceback))
    
    # Print to console
    print("\n" + "="*80)
    print("CRASH - Ứng dụng đã gặp lỗi nghiêm trọng!")
    print("="*80)
    print(f"Loại lỗi: {exc_type.__name__}")
    print(f"Thông báo: {error_msg}")
    print("\nTraceback chi tiết:")
    print(traceback_str)
    print("="*80)
    
    # Save error log to file
    try:
        import datetime
        log_file = "crash_log.txt"
        with open(log_file, 'a', encoding='utf-8') as f:
            f.write("\n" + "="*80 + "\n")
            f.write(f"Crash Time: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Error Type: {exc_type.__name__}\n")
            f.write(f"Error Message: {error_msg}\n")
            f.write("\nFull Traceback:\n")
            f.write(traceback_str)
            f.write("="*80 + "\n\n")
        print(f"\nĐã lưu log vào: {log_file}")
    except:
        pass
    
    # Try to show QMessageBox if QApplication exists
    try:
        app = QApplication.instance()
        if app is not None:
            # Create detailed error message for popup
            detailed_msg = (
                f"ỨNG DỤNG ĐÃ CRASH!\n\n"
                f"Loại lỗi: {exc_type.__name__}\n"
                f"Thông báo: {error_msg}\n\n"
                f"Chi tiết kỹ thuật:\n"
                f"{traceback_str[:500]}..." if len(traceback_str) > 500 else traceback_str
            )
            
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Icon.Critical)
            msg.setWindowTitle("Lỗi Nghiêm Trọng - Ứng Dụng Đã Crash")
            msg.setText("Ứng dụng đã gặp lỗi và cần đóng.\n\nVui lòng xem chi tiết bên dưới để fix bug.")
            msg.setDetailedText(traceback_str)
            msg.setStandardButtons(QMessageBox.StandardButton.Ok)
            msg.setDefaultButton(QMessageBox.StandardButton.Ok)
            msg.exec()
        else:
            # If no QApplication, create one just for the error dialog
            app = QApplication([])
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Icon.Critical)
            msg.setWindowTitle("Lỗi Nghiêm Trọng - Ứng Dụng Đã Crash")
            msg.setText(
                f"Ứng dụng đã crash ngay khi khởi động!\n\n"
                f"Loại lỗi: {exc_type.__name__}\n"
                f"Thông báo: {error_msg}\n\n"
                f"Xem console để biết thêm chi tiết."
            )
            msg.setDetailedText(traceback_str[:2000])  # Limit detail text length
            msg.setStandardButtons(QMessageBox.StandardButton.Ok)
            msg.exec()
            sys.exit(1)
    except Exception as e:
        # If even showing error dialog fails, just print
        print(f"\nKhông thể hiển thị popup lỗi: {e}")
        print("\nNhấn Enter để thoát...")
        try:
            input()
        except:
            pass
    
    # Exit with error code
    sys.exit(1)


def main():
    try:
        import sys
        import traceback
        
        # Set global exception handler
        sys.excepthook = exception_hook
        
        # Suppress transformers loading messages if needed
        import warnings
        warnings.filterwarnings("ignore", category=UserWarning)
        
        print("=== Khởi động ứng dụng OCR ===")
        print("Lưu ý: Auto-load đã TẮT mặc định. Click nút 'Sử dụng AI này' để load thủ công.")
        print("=" * 50)
        
        app = QApplication(sys.argv)
        
        try:
            print("Đang tạo UI...")
            window = Qwen3VLApp()
            window.show()
            print("UI đã được tạo thành công!")
            print("Auto-load:", "BẬT" if window.auto_load_enabled else "TẮT")
        except Exception as e:
            # If window creation fails, exception hook will handle it
            raise
        
        # Run application event loop
        try:
            print("Ứng dụng đang chạy...")
            sys.exit(app.exec())
        except KeyboardInterrupt:
            print("\nỨng dụng đã bị người dùng dừng.")
            return 0
        except SystemExit as e:
            raise e
        except Exception as e:
            # Re-raise to let exception hook handle it
            raise
    except SystemExit as e:
        # Re-raise system exit
        raise e
    except Exception as e:
        # This should be caught by exception_hook, but just in case
        raise


if __name__ == "__main__":
    main()

